from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from pipeline import ResumePipeline


@pytest.mark.integration
def test_real_pdf_runs_end_to_end_and_exports_identical_json(tmp_path: Path) -> None:
    fixture = Path(__file__).parent / "fixtures" / "resume-accounting.pdf"
    output = tmp_path / "accounting-report.json"
    result = ResumePipeline().analyze(str(fixture), output_path=output)

    assert result["schema_version"] == "2.1.0"
    assert result["ats"]["status"] in {"complete", "partial"}
    assert result["document"]["pages"] >= 1
    assert result["entities"]["skills"]
    experience = result["entities"]["experience"]
    bullets = [
        value for item in experience for value in (*item["responsibilities"], *item["achievements"])
    ]
    assert len(experience) == 4
    assert len(bullets) == 15
    assert {
        "Reconciled bank and credit card accounts for up to 5 clients "
        "simultaneously while maintaining a high degree of accuracy",
        "Increased efficiency and alleviated work loads by creating a new "
        "Excel financial recording system, saved reporting time by 15%",
        "Implemented effective sales strategies which led to 120% achievement "
        "of seasonal sales targets established for the summer of 2017",
        "Successfully closed the store's biggest sale in one day, amounting "
        "to $1,200, by providing excellent customer service using up-selling "
        "and cross-selling strategies",
        "Performed administrative duties including faxing, filing and managing "
        "inbound calls and emails, making sure office operations ran smoothly",
    } <= set(bullets)
    assert not any("\u00ad" in value for value in bullets)
    assert "Sales Associate" in {item["job_title"] for item in experience}
    assert result["entities"]["projects"] == []
    assert result["module_status"]["target_role"]["status"] == "complete"
    assert result["module_status"]["recommendations"]["status"] in {"complete", "fallback"}
    assert json.loads(output.read_text(encoding="utf-8")) == result


@pytest.mark.integration
def test_real_docx_runs_end_to_end_without_optional_models(tmp_path: Path) -> None:
    source = tmp_path / "backend-resume.docx"
    document = Document()
    document.add_heading("Jane Doe", 0)
    document.add_paragraph("jane@example.com | +1 555 0100")
    document.add_heading("Summary", level=1)
    document.add_paragraph(
        "Backend engineer building reliable production APIs and data services with Python."
    )
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python, FastAPI, SQL, Docker, Git, REST API")
    document.add_heading("Experience", level=1)
    document.add_paragraph("Backend Engineer | Example Labs | 2022 - Present")
    document.add_paragraph("Built Python APIs for internal services.")
    document.add_heading("Education", level=1)
    document.add_paragraph("Bachelor of Science in Computer Science | Example University | 2021")
    document.save(source)

    result = ResumePipeline().analyze(str(source))

    assert result["document"]["extension"] == ".docx"
    assert result["extraction"]["engine"] == "docx_ooxml"
    assert result["module_status"]["extraction"]["status"] in {"complete", "degraded"}
    assert result["module_status"]["target_role"]["status"] == "complete"
    assert result["errors"] == []
