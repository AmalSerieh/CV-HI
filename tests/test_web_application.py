from __future__ import annotations

import io
import json
import re
import threading
import time
import zipfile
from copy import deepcopy
from datetime import datetime, timedelta, timezone
from pathlib import Path

import fitz
import pytest
from docx import Document
from fastapi import Request
from fastapi.testclient import TestClient

from pipeline import PipelineConfig, ResumePipeline
from resume_analyzer.ats import ATSAnalyzer
from resume_analyzer.export import FinalResumeBuilder
from resume_analyzer.extraction.data_quality import CanonicalDataQualityAnalyzer
from resume_analyzer.schemas import PipelineReport
from resume_analyzer.web.app import create_app
from resume_analyzer.web.app import main as web_main
from resume_analyzer.web.config import WebSettings
from tests.report_fixtures import make_report
from tests.test_resume_export import reviewable_report


class ReportBackend:
    def __init__(self, report):
        self.report = report

    def extract(self, _file_path):
        return self.report

    def extract_text(self, _text, *, document_name):
        return self.report


class RecordingFactory:
    def __init__(self, report):
        self.report = report
        self.calls = 0
        self.configs = []
        self.paths = []
        self.job_descriptions = []

    def __call__(self, config):
        self.configs.append(config)
        outer = self

        class PipelineStub:
            def analyze(self, file_path, *, job_description=None):
                outer.calls += 1
                outer.paths.append(Path(file_path))
                outer.job_descriptions.append(job_description)
                return deepcopy(outer.report)

        return PipelineStub()


class FailingFactory:
    def __call__(self, _config):
        class PipelineStub:
            def analyze(self, _file_path, *, job_description=None):
                raise RuntimeError(r"private failure at C:\Users\someone\resume.pdf")

        return PipelineStub()


class BlockingFactory(RecordingFactory):
    def __init__(self, report):
        super().__init__(report)
        self.started = threading.Event()
        self.release = threading.Event()

    def __call__(self, config):
        self.configs.append(config)
        outer = self

        class PipelineStub:
            def analyze(self, file_path, *, job_description=None):
                outer.calls += 1
                outer.paths.append(Path(file_path))
                outer.started.set()
                outer.release.wait(timeout=5)
                return deepcopy(outer.report)

        return PipelineStub()


@pytest.fixture(scope="session")
def rich_report():
    base = make_report()
    config = PipelineConfig(
        enable_ocr=False,
        integrate_target_role=True,
        enable_recommendations=True,
        enable_ats=True,
        enable_rewrites=True,
        ai_provider="none",
    )
    return ResumePipeline(config=config, extraction_backend=ReportBackend(base)).analyze(
        "synthetic.pdf",
        job_description="Python API engineer with SQL and PostgreSQL experience",
    )


@pytest.fixture
def settings(tmp_path):
    return WebSettings(
        temp_dir=tmp_path,
        result_ttl_minutes=1,
        max_upload_mb=1,
        max_pages=2,
        max_extracted_chars=20_000,
        max_job_description_chars=1_000,
        max_concurrent_analyses=2,
        max_docx_uncompressed_mb=5,
        max_docx_files=200,
    )


def test_app_creates_configured_runtime_directories(tmp_path, rich_report) -> None:
    temp_dir = tmp_path / "runtime" / "temp"
    output_dir = tmp_path / "runtime" / "outputs"
    selected = WebSettings(temp_dir=temp_dir, output_dir=output_dir)
    app = create_app(selected, pipeline_factory=RecordingFactory(rich_report))
    assert app.state.settings is selected
    assert temp_dir.is_dir()
    assert output_dir.is_dir()


def test_web_main_uses_reload_setting_and_rejects_an_existing_listener(
    monkeypatch, settings
) -> None:
    calls = []
    selected = WebSettings(**{**settings.__dict__, "reload": True})
    monkeypatch.setattr(
        "resume_analyzer.web.app.WebSettings.from_env",
        lambda: selected,
    )
    monkeypatch.setattr("resume_analyzer.web.app._port_is_in_use", lambda *_args: False)
    monkeypatch.setattr(
        "resume_analyzer.web.app.uvicorn.run", lambda *args, **kwargs: calls.append((args, kwargs))
    )
    web_main()
    assert calls[0][1]["reload"] is True
    assert calls[0][1]["reload_dirs"] == [str(Path("resume_analyzer").resolve())]

    monkeypatch.setattr("resume_analyzer.web.app._port_is_in_use", lambda *_args: True)
    with pytest.raises(SystemExit, match="Ctrl\\+C"):
        web_main()


@pytest.fixture
def web_client(settings, rich_report):
    factory = RecordingFactory(rich_report)
    with TestClient(create_app(settings, pipeline_factory=factory)) as client:
        yield client, factory


def pdf_bytes(text="Jane Doe Python resume", pages=1):
    document = fitz.open()
    for index in range(pages):
        page = document.new_page()
        page.insert_text((72, 72), f"{text} {index + 1}")
    value = document.tobytes()
    document.close()
    return value


def docx_bytes(text="Jane Doe Python resume"):
    stream = io.BytesIO()
    document = Document()
    document.add_heading("Resume", 0)
    document.add_paragraph(text)
    document.save(stream)
    return stream.getvalue()


def resume_file(value=None, name="resume.pdf", mime="application/pdf"):
    return {"resume": (name, value if value is not None else pdf_bytes(), mime)}


def submit(client, *, files=None, data=None):
    return client.post(
        "/api/analyses",
        files=files or resume_file(),
        data=data or {},
    )


def wait_for_terminal(client, status_url, timeout=3):
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        response = client.get(status_url)
        if response.status_code == 200 and response.json()["status"] in {
            "completed",
            "failed",
        }:
            return response.json()
        time.sleep(0.01)
    raise AssertionError("analysis did not reach a terminal state")


def completed_submission(client, **kwargs):
    response = submit(client, **kwargs)
    assert response.status_code == 202
    payload = response.json()
    status = wait_for_terminal(client, payload["status_url"])
    assert status["status"] == "completed"
    return payload


def test_application_openapi_starts(web_client):
    client, _ = web_client
    assert client.get("/openapi.json").status_code == 200


def test_landing_page(web_client):
    client, _ = web_client
    response = client.get("/")
    assert response.status_code == 200
    assert "Resume Intelligence Platform" in response.text
    assert "Runs locally" in response.text


def test_analysis_form(web_client):
    client, _ = web_client
    response = client.get("/analyze")
    assert response.status_code == 200
    assert 'id="analysis-form"' in response.text
    assert 'accept=".pdf,.docx' in response.text


@pytest.mark.parametrize(
    "path,expected",
    [
        ("/static/css/app.css", "--app-navy"),
        ("/static/css/rtl.css", 'html[dir="rtl"]'),
        ("/static/css/print.css", "@page"),
        ("/static/js/app.js", '"use strict"'),
        ("/static/js/upload.js", "FormData"),
        ("/static/js/progress.js", "setTimeout"),
        ("/static/js/results.js", "clipboard"),
    ],
)
def test_local_static_assets(web_client, path, expected):
    client, _ = web_client
    response = client.get(path)
    assert response.status_code == 200
    assert expected in response.text


@pytest.mark.parametrize(
    "path,marker",
    [
        ("/static/vendor/bootstrap/bootstrap.min.css", "Bootstrap"),
        ("/static/vendor/bootstrap/bootstrap.bundle.min.js", "Bootstrap"),
    ],
)
def test_local_bootstrap_assets(web_client, path, marker):
    client, _ = web_client
    response = client.get(path)
    assert response.status_code == 200
    assert marker in response.text


def test_security_headers(web_client):
    client, _ = web_client
    response = client.get("/")
    assert response.headers["x-frame-options"] == "DENY"
    assert "object-src 'none'" in response.headers["content-security-policy"]


def test_health_endpoint(web_client):
    client, _ = web_client
    response = client.get("/api/health")
    payload = response.json()
    assert payload["status"] == "ok"
    assert payload["schema_version"] == "2.1.0"
    assert payload["build"]["status"] == "current"
    assert payload["build"]["restart_required"] is False
    assert response.headers["x-resume-source-stale"] == "false"


def test_stale_backend_blocks_new_analysis_and_explains_restart(settings, rich_report) -> None:
    app = create_app(settings, pipeline_factory=RecordingFactory(rich_report))
    app.state.source_fingerprint_provider = lambda: "f" * 64
    with TestClient(app) as client:
        health = client.get("/api/health")
        assert health.json()["build"]["restart_required"] is True
        assert health.headers["x-resume-source-stale"] == "true"
        page = client.get("/analyze")
        assert "Backend restart required" in page.text
        response = submit(client)
        assert response.status_code == 503
        assert response.json()["error"]["code"] == "backend_restart_required"


def test_system_capabilities_endpoint(web_client):
    client, _ = web_client
    response = client.get("/api/system")
    assert response.status_code == 200
    assert response.json()["frontend"]["bootstrap_local"] is True


def test_models_endpoint_is_privacy_safe(web_client):
    client, _ = web_client
    serialized = client.get("/api/models").text
    assert "fallback_available" in serialized
    assert "C:\\Users\\" not in serialized
    assert '"endpoint"' not in serialized


def test_pdf_upload_and_analysis_creation(web_client):
    client, _ = web_client
    response = submit(client)
    assert response.status_code == 202
    assert response.json()["analysis_id"] in response.json()["status_url"]


def test_docx_upload(web_client):
    client, _ = web_client
    payload = completed_submission(
        client,
        files={
            "resume": (
                "resume.docx",
                docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert client.get(payload["result_url"]).json()["document"]["name"] == "resume.docx"


@pytest.mark.parametrize("name", ["resume.exe", "resume.txt", "resume.docm", "resume.PDF.exe"])
def test_unsupported_or_macro_extension(web_client, name):
    client, _ = web_client
    response = submit(client, files=resume_file(b"payload", name, "application/octet-stream"))
    assert response.status_code == 400
    assert response.json()["error"]["code"] in {
        "unsupported_file_type",
        "macro_format_rejected",
    }


def test_empty_file_rejected(web_client):
    client, _ = web_client
    response = submit(client, files=resume_file(b""))
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "empty_file"


def test_oversized_file_rejected(web_client):
    client, _ = web_client
    response = submit(client, files=resume_file(b"%PDF-" + b"x" * 1_000_001))
    assert response.status_code == 413
    assert response.json()["error"]["code"] == "file_too_large"


def test_invalid_mime_type_rejected(web_client):
    client, _ = web_client
    response = submit(client, files=resume_file(pdf_bytes(), mime="text/plain"))
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "invalid_mime_type"


def test_malformed_pdf_rejected(web_client):
    client, _ = web_client
    response = submit(client, files=resume_file(b"%PDF-not-a-document"))
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "corrupt_document"


def test_malformed_docx_rejected(web_client):
    client, _ = web_client
    response = submit(
        client,
        files={
            "resume": (
                "resume.docx",
                b"PK\x03\x04not-a-docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "corrupt_document"


def test_macro_part_inside_docx_rejected(web_client):
    client, _ = web_client
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "document")
        archive.writestr("word/vbaProject.bin", b"macro")
    response = submit(
        client,
        files={
            "resume": (
                "resume.docx",
                stream.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "macro_format_rejected"


def test_docx_zip_bomb_ratio_rejected(web_client):
    client, _ = web_client
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "a" * 2_000_000)
    response = submit(
        client,
        files={
            "resume": (
                "resume.docx",
                stream.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "unsafe_archive"


def test_pdf_page_limit_rejected(web_client):
    client, _ = web_client
    response = submit(client, files=resume_file(pdf_bytes(pages=3)))
    assert response.status_code == 413
    assert response.json()["error"]["code"] == "too_many_pages"


@pytest.mark.parametrize("name", ["../resume.pdf", r"..\resume.pdf", "/tmp/resume.pdf"])
def test_path_traversal_filename_rejected(web_client, name):
    client, _ = web_client
    response = submit(client, files=resume_file(pdf_bytes(), name=name))
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "invalid_filename"


def test_job_description_text_reaches_pipeline(web_client):
    client, factory = web_client
    completed_submission(client, data={"job_description_text": "Python SQL role"})
    assert factory.job_descriptions[-1] == "Python SQL role"


def test_job_description_txt_upload_reaches_pipeline(web_client):
    client, factory = web_client
    files = resume_file()
    files["job_description_file"] = ("role.txt", "Arabic العربية Python", "text/plain")
    completed_submission(client, files=files)
    assert factory.job_descriptions[-1] == "Arabic العربية Python"


def test_invalid_job_description_extension(web_client):
    client, _ = web_client
    files = resume_file()
    files["job_description_file"] = ("role.html", "<script>", "text/html")
    response = submit(client, files=files)
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "invalid_job_description_file"


def test_oversized_job_description(settings, rich_report):
    selected = WebSettings(**{**settings.__dict__, "max_job_description_chars": 5})
    factory = RecordingFactory(rich_report)
    with TestClient(create_app(selected, pipeline_factory=factory)) as client:
        response = submit(client, data={"job_description_text": "sixsix"})
    assert response.status_code == 413
    assert response.json()["error"]["code"] == "job_description_too_large"


def test_status_polling_and_stage_contract(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    status = client.get(payload["status_url"]).json()
    assert status["stage"] == "completed"
    assert status["completed_stages"] == [
        "uploading",
        "validating_document",
        "running_pipeline",
        "validating_final_report",
        "completed",
    ]


def test_result_retrieval(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    result = client.get(payload["result_url"])
    assert result.status_code == 200
    assert result.json()["schema_version"] == "2.1.0"


def test_result_page_rendering(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    page = client.get(payload["page_url"])
    assert page.status_code == 200
    assert "Canonical report" in page.text
    assert "Candidate information" in page.text


def test_json_download(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    response = client.get(payload["result_url"].replace("/result", "/download"))
    assert response.status_code == 200
    assert response.headers["content-disposition"].startswith("attachment;")
    assert response.json()["schema_version"] == "2.1.0"


def test_runtime_and_download_json_are_identical(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    runtime = client.get(payload["result_url"]).json()
    downloaded = client.get(payload["result_url"].replace("/result", "/download")).json()
    assert runtime == downloaded


def test_temporary_resume_is_deleted_after_completion(web_client):
    client, factory = web_client
    completed_submission(client)
    assert factory.paths
    assert not factory.paths[-1].parent.exists()


def test_explicit_analysis_delete(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    assert client.delete(payload["status_url"]).status_code == 204
    assert client.get(payload["status_url"]).status_code == 404


def test_ttl_cleanup_removes_expired_record(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    identifier = payload["analysis_id"]
    record = client.app.state.job_store.get(__import__("uuid").UUID(identifier))
    record.updated_at = datetime.now(timezone.utc) - timedelta(minutes=2)
    assert client.get(payload["status_url"]).status_code == 404
    review_url = f"/api/analyses/{identifier}/resume-review"
    assert client.get(review_url).status_code == 404
    assert client.get(f"/api/analyses/{identifier}/final-resume").status_code == 404
    assert (
        client.post(
            f"/api/analyses/{identifier}/download-docx",
            json={"template_id": "template-1"},
        ).status_code
        == 404
    )


def test_unknown_analysis_id(web_client):
    client, _ = web_client
    assert client.get("/api/analyses/00000000-0000-4000-8000-000000000000").status_code == 404


def test_invalid_uuid(web_client):
    client, _ = web_client
    assert client.get("/api/analyses/not-a-uuid").status_code == 422
    assert client.get("/results/not-a-uuid").status_code == 400


def test_result_is_conflict_until_ready(settings, rich_report):
    factory = BlockingFactory(rich_report)
    with TestClient(create_app(settings, pipeline_factory=factory)) as client:
        response = submit(client)
        assert factory.started.wait(timeout=1)
        assert client.get(response.json()["result_url"]).status_code == 409
        identifier = response.json()["analysis_id"]
        assert client.get(f"/api/analyses/{identifier}/resume-review").status_code == 409
        assert client.get(f"/api/analyses/{identifier}/final-resume").status_code == 409
        assert (
            client.post(
                f"/api/analyses/{identifier}/download-docx",
                json={"template_id": "template-1"},
            ).status_code
            == 409
        )
        factory.release.set()


def test_progress_page_lists_truthful_stages(settings, rich_report):
    factory = BlockingFactory(rich_report)
    with TestClient(create_app(settings, pipeline_factory=factory)) as client:
        response = submit(client)
        assert factory.started.wait(timeout=1)
        page = client.get(response.json()["page_url"])
        assert "does not invent percentages" in page.text
        assert "Extracting text" in page.text
        assert "Validating final report" in page.text
        factory.release.set()


def test_concurrent_analysis_limit(settings, rich_report):
    selected = WebSettings(**{**settings.__dict__, "max_concurrent_analyses": 1})
    factory = BlockingFactory(rich_report)
    with TestClient(create_app(selected, pipeline_factory=factory)) as client:
        first = submit(client)
        assert first.status_code == 202
        assert factory.started.wait(timeout=1)
        second = submit(client)
        assert second.status_code == 429
        assert second.json()["error"]["code"] == "analysis_capacity_reached"
        factory.release.set()


def test_pipeline_called_only_once(web_client):
    client, factory = web_client
    payload = completed_submission(client)
    for _ in range(3):
        assert client.get(payload["status_url"]).status_code == 200
    assert client.get(payload["result_url"]).status_code == 200
    assert client.get(payload["page_url"]).status_code == 200
    assert client.get(payload["result_url"].replace("/result", "/download")).status_code == 200
    assert factory.calls == 1


def test_ui_options_map_to_pipeline_config(web_client):
    client, factory = web_client
    completed_submission(
        client,
        data={
            "enable_target_role": "false",
            "enable_recommendations": "false",
            "enable_ats": "false",
            "enable_rewrites": "true",
            "enable_ocr": "false",
            "output_language": "ar",
            "rewrite_summary": "true",
            "rewrite_experience": "false",
            "rewrite_skills": "false",
            "bullet_rewrite_mode": "specific",
            "bullet_rewrite_selection": "2, 4",
        },
    )
    config = factory.configs[-1]
    assert config.integrate_target_role is False
    assert config.enable_recommendations is False
    assert config.enable_ats is False
    assert config.enable_rewrites is True
    assert config.enable_ocr is False
    assert config.rewrite_language == "ar"
    assert config.rewrite_sections == ("summary",)
    assert config.rewrite_bullet_selection == (1, 3)
    assert config.rewrite_all_bullets is False


def test_rejected_summary_ui_has_no_duplicate_proposed_card(settings, rich_report):
    report = deepcopy(rich_report)
    original = report["entities"]["summary"]
    report["rewrites"]["summary"].update(
        status="rejected",
        original=original,
        improved=None,
        warnings=["The local model response was incomplete."],
    )
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])
    summary_html = page.text.split("Summary rewrite", 1)[1].split("Experience bullets", 1)[0]
    assert "No valid rewrite was generated." in summary_html
    assert '<h3 class="h6">Proposed</h3>' not in summary_html
    assert summary_html.count(original) == 1


def test_rejected_bullet_ui_displays_the_validation_reason(settings, rich_report):
    report = deepcopy(rich_report)
    bullet = report["rewrites"]["experience_bullets"][0]
    bullet.update(
        status="rejected",
        improved=None,
        warnings=["Unsupported factual claim: invented certification."],
    )
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])

    bullet_html = page.text.split("Experience bullets", 1)[1].split("Skills organization", 1)[0]
    assert "No valid rewrite was generated." in bullet_html
    assert "Unsupported factual claim: invented certification." in bullet_html


def test_analysis_form_exposes_bounded_bullet_selection(web_client):
    client, _ = web_client
    page = client.get("/analyze")
    assert 'name="bullet_rewrite_mode" value="first"' in page.text
    assert 'name="bullet_rewrite_mode" value="specific"' in page.text
    assert 'name="bullet_rewrite_mode" value="all"' in page.text
    assert "A hard safety limit of 20 still applies." in page.text


def test_analysis_form_explains_model_responsibilities_and_uses_configured_default(
    web_client, monkeypatch
):
    monkeypatch.setattr(
        "resume_analyzer.web.routes.pages.model_status",
        lambda **_kwargs: {
            "configured_provider": "ollama",
            "configured_model": "gemma3:4b",
            "ollama": {
                "reachable": True,
                "available_models": ["gemma3:4b", "qwen3:4b"],
            },
            "transformers": {"installed": False},
        },
    )
    client, _ = web_client
    page = client.get("/analyze")
    assert 'value="ollama" selected' in page.text
    assert 'value="gemma3:4b"' in page.text
    assert "The model affects recommendations and rewrite proposals only." in page.text
    assert "The percentage is a weighted evidence-match score" in page.text
    assert "rules only, no generated rewrites" in page.text


def test_analysis_form_survives_local_model_discovery_failure(web_client, monkeypatch):
    monkeypatch.setattr(
        "resume_analyzer.web.routes.pages.model_status",
        lambda **_kwargs: (_ for _ in ()).throw(OSError("synthetic discovery failure")),
    )
    client, _ = web_client
    page = client.get("/analyze")
    assert page.status_code == 200
    assert 'value="none" selected' in page.text
    assert "Ollama is not currently reachable." in page.text


def test_analysis_template_is_backward_compatible_without_models_context(settings, rich_report):
    app = create_app(settings, pipeline_factory=RecordingFactory(rich_report))

    @app.get("/test/analyze-without-model-context")
    def legacy_analyze_page(request: Request):
        return request.app.state.templates.TemplateResponse(
            request=request,
            name="analyze.html",
            context={
                "page_title": "Analyze a resume",
                "direction": "ltr",
                "settings": request.app.state.settings,
            },
        )

    with TestClient(app) as client:
        page = client.get("/test/analyze-without-model-context")
    assert page.status_code == 200
    assert 'id="analysis-form"' in page.text
    assert 'value="none" selected' in page.text


def test_successful_summary_is_not_duplicated_in_page_markup(settings, rich_report):
    report = deepcopy(rich_report)
    proposal = "A unique validated proposal rendered exactly once."
    report["rewrites"]["status"] = "complete"
    report["rewrites"]["summary"].update(
        status="improved",
        original=report["entities"]["summary"],
        improved=proposal,
        warnings=[],
        requires_review=True,
        generated_from_evidence=True,
    )
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])
    summary_html = page.text.split("Summary rewrite", 1)[1].split("Experience bullets", 1)[0]
    assert summary_html.count(proposal) == 1
    assert 'id="rewrite-summary"' in summary_html


def test_results_hide_raw_evidence_ids_and_load_json_on_demand(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    report = client.get(payload["result_url"]).json()
    page = client.get(payload["page_url"])
    assert report["evidence"][0]["id"] not in page.text
    assert "internal IDs are intentionally not displayed here" in page.text
    assert "Developer-oriented technical data" in page.text
    assert '<pre id="raw-json" class="json-view mt-3 d-none"></pre>' in page.text


def test_results_explain_rule_scores_and_not_run_job_match(settings, rich_report):
    report = deepcopy(rich_report)
    report["ats"]["job_match"].update(
        status="not_run",
        match_score=None,
        matched_keywords=[],
        missing_keywords=[],
    )
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])
    assert "it is not a hiring probability or seniority rating" in page.text
    assert "How the score was calculated" in page.text
    assert "Text Extractability</span><strong>15 / 15" in page.text
    assert "Not calculated." in page.text
    assert "Method: evidence-based rules." in page.text


def test_job_match_can_be_disabled(web_client):
    client, factory = web_client
    completed_submission(
        client,
        data={"job_description_text": "private role", "enable_job_match": "false"},
    )
    assert factory.job_descriptions[-1] is None


def test_provider_failure_is_sanitized(settings):
    with TestClient(create_app(settings, pipeline_factory=FailingFactory())) as client:
        response = submit(client)
        status = wait_for_terminal(client, response.json()["status_url"])
    assert status["status"] == "failed"
    assert "C:\\Users" not in json.dumps(status)
    assert status["error"]["code"] == "analysis_failed"


def test_no_model_recommendation_fallback_renders(web_client):
    client, _ = web_client
    payload = completed_submission(client)
    result = client.get(payload["result_url"]).json()
    assert result["module_status"]["recommendations"]["status"] == "fallback"


def test_hybrid_recommendation_label_is_honest(settings, rich_report):
    report = deepcopy(rich_report)
    summary_evidence_id = next(
        item["id"] for item in report["evidence"] if item["field_path"] == "entities.summary"
    )
    report["module_status"]["recommendations"].update(
        status="complete",
        provider="ollama",
        model="gemma3:4b",
        detail=None,
    )
    report["recommendations"] = [
        {
            "id": "rec-hybrid-summary",
            "area": "summary",
            "severity": "low",
            "confidence": 0.95,
            "title": "Improve summary clarity",
            "problem": "The cited summary can be clearer and more concise.",
            "suggestion": "Refine the cited summary while preserving every supported fact.",
            "evidence_ids": [summary_evidence_id],
            "source": "hybrid",
            "conditional": False,
        }
    ]
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])
    assert "Method: deterministic evidence focus with validated local-model response." in page.text
    assert "Evidence ranked + local model" in page.text
    assert "Method: local AI." not in page.text


def test_tesseract_unavailable_is_nonfatal(web_client, monkeypatch):
    client, _ = web_client
    monkeypatch.setattr(
        "resume_analyzer.diagnostics.health.tesseract_status",
        lambda: {"installed": False, "languages": [], "error": "not found"},
    )
    response = client.get("/api/system")
    assert response.status_code == 200
    assert response.json()["ocr"]["installed"] is False


def test_arabic_result_uses_rtl(settings, rich_report):
    report = deepcopy(rich_report)
    report["target_role"]["language"] = "ar"
    report["entities"]["contact"]["name"] = "ليلى أحمد"
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        payload = completed_submission(client)
        page = client.get(payload["page_url"])
    assert '<html lang="ar" dir="rtl">' in page.text
    assert "ليلى أحمد" in page.text


def test_english_result_uses_ltr(web_client):
    client, _ = web_client
    page = client.get(completed_submission(client)["page_url"])
    assert 'dir="ltr"' in page.text


def test_mixed_language_content_renders_utf8(settings, rich_report):
    report = deepcopy(rich_report)
    report["entities"]["summary"] = "مهندس Python builds APIs"
    report["rewrites"]["summary"]["original"] = "مهندس Python builds APIs"
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])
    assert "مهندس Python builds APIs" in page.text


@pytest.mark.parametrize(
    "needle",
    [
        "ATS compatibility",
        "Parsing integrity",
        "Extraction confidence",
        "Layout risk",
        "Contact readability",
        "Job-description match",
        "Recommendations",
        "Review-only proposals",
        "Target-role evidence match",
        "Evidence and audit trail",
        "Warnings and errors",
        "Canonical JSON",
    ],
)
def test_results_dashboard_renders_every_canonical_area(web_client, needle):
    client, _ = web_client
    page = client.get(completed_submission(client)["page_url"])
    assert needle in page.text


def test_unreliable_contact_and_image_contact_are_rendered_conservatively(settings, rich_report):
    report = deepcopy(rich_report)
    report["entities"]["contact"]["phone"] = None
    report["extraction"]["visual_metadata"]["possible_image_only_contact"] = True
    report["extraction"]["visual_metadata"]["contact_readability"] = "partially_readable"
    report["data_quality"]["contact_readability"] = "partially_readable"
    report["data_quality"]["fields_requiring_review"] = ["entities.contact.phone"]
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])

    assert "No reliable phone number was extracted." in page.text
    assert "Contact information may be embedded in an image." in page.text
    assert "Needs review" in page.text


def test_ocr_contact_does_not_render_false_selectable_text_strength(settings, rich_report):
    report = deepcopy(rich_report)
    report["entities"]["contact"]["source_types"].update(email="ocr", phone="ocr")
    report["extraction"]["ocr_used"] = True
    report["extraction"]["ocr_usage"] = {
        "used": True,
        "scope": "contact_header",
        "pages": [1],
        "fields": ["email", "phone"],
    }
    report["extraction"]["visual_metadata"].update(
        contact_ocr_used=True,
        contact_ocr_status="complete",
    )
    canonical = PipelineReport.model_validate(report)
    report["ats"] = ATSAnalyzer().analyze(canonical).model_dump(mode="json")
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])

    assert "Required contact details are selectable text" not in page.text
    assert "OCR-recovered contact details need verification" in page.text


def test_contact_ocr_scope_is_rendered_without_global_warning(settings, rich_report):
    report = deepcopy(rich_report)
    report["entities"]["contact"]["source_types"].update(email="ocr", phone="ocr")
    report["extraction"]["ocr_used"] = True
    report["extraction"]["ocr_usage"] = {
        "used": True,
        "scope": "contact_header",
        "pages": [1],
        "fields": ["email", "phone"],
    }
    report["extraction"]["visual_metadata"].update(
        contact_ocr_used=True,
        contact_ocr_status="complete",
    )
    canonical = PipelineReport.model_validate(report)
    report["ats"] = ATSAnalyzer().analyze(canonical).model_dump(mode="json")
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])

    assert "OCR scope:" in page.text
    assert "contact header" in page.text
    assert "Some or all text was recovered with OCR" not in page.text


def test_experience_parsing_and_content_review_render_separately(settings, rich_report):
    report = deepcopy(rich_report)
    experience = report["entities"]["experience"][0]
    experience.update(
        needs_review=True,
        parsing_needs_review=False,
        content_needs_review=True,
        review_reasons=["truncated_experience_bullets: An experience bullet may be incomplete."],
    )
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])

    assert "Parsing verified" in page.text
    assert "Content needs review" in page.text
    assert "Content review notes" in page.text
    assert "truncated experience bullets" in page.text


def test_parsing_integrity_breakdown_is_rendered(settings, rich_report):
    report = deepcopy(rich_report)
    canonical = PipelineReport.model_validate(report)
    report["data_quality"] = (
        CanonicalDataQualityAnalyzer().analyze(canonical).model_dump(mode="json")
    )
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])

    assert "How parsing integrity was calculated" in page.text
    assert "Weighted subtotal:" in page.text
    assert "Contact Integrity" in page.text
    assert "Final parsing integrity:" in page.text


def test_absolute_paths_removed_from_api(settings, rich_report):
    report = deepcopy(rich_report)
    report["document"]["path"] = r"C:\Users\private\resume.pdf"
    report["warnings"].append(
        {
            "stage": "test",
            "code": "path_test",
            "message": r"failed at C:\Users\private\resume.pdf",
            "recoverable": True,
        }
    )
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        payload = completed_submission(client)
        serialized = client.get(payload["result_url"]).text
    assert "C:\\Users" not in serialized
    assert '"path":null' in serialized


def test_html_and_script_injection_are_escaped(settings, rich_report):
    report = deepcopy(rich_report)
    report["entities"]["contact"]["name"] = '<script>alert("x")</script>'
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        page = client.get(completed_submission(client)["page_url"])
    assert '<script>alert("x")</script>' not in page.text
    assert "&lt;script&gt;alert" in page.text


def test_frontend_does_not_use_untrusted_inner_html():
    static_dir = Path("resume_analyzer/web/static/js")
    for path in static_dir.glob("*.js"):
        assert "innerHTML" not in path.read_text(encoding="utf-8")


def test_download_filename_never_uses_uploaded_name(web_client):
    client, _ = web_client
    payload = completed_submission(client, files=resume_file(name="candidate-private.pdf"))
    response = client.get(payload["result_url"].replace("/result", "/download"))
    assert "candidate-private" not in response.headers["content-disposition"]
    assert payload["analysis_id"] in response.headers["content-disposition"]


def test_resume_review_api_persists_decisions_and_keeps_analysis_immutable(settings):
    report = reviewable_report().model_dump(mode="json")
    factory = RecordingFactory(report)
    with TestClient(create_app(settings, pipeline_factory=factory)) as client:
        payload = completed_submission(client)
        result_before = client.get(payload["result_url"]).json()
        review_url = f"/api/analyses/{payload['analysis_id']}/resume-review"
        initial = client.get(review_url)
        assert initial.status_code == 200
        assert initial.json()["summary"]["decision"] == "pending"
        bullet_id = initial.json()["experience_bullets"][0]["id"]

        updated = client.patch(
            review_url,
            json={
                "summary": "accepted",
                "skills": "accepted",
                "experience_bullets": {bullet_id: "accepted"},
            },
        )
        assert updated.status_code == 200
        assert updated.json()["summary"]["final"] == ("Python engineer delivering reliable APIs.")
        assert updated.json()["experience_bullets"][0]["final"] == ("Built reliable Python APIs.")
        assert client.get(review_url).json()["summary"]["decision"] == "accepted"
        assert client.get(payload["result_url"]).json() == result_before


def test_review_update_is_transactional_for_unknown_item(settings):
    report = reviewable_report().model_dump(mode="json")
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        payload = completed_submission(client)
        review_url = f"/api/analyses/{payload['analysis_id']}/resume-review"
        response = client.patch(
            review_url,
            json={
                "summary": "accepted",
                "experience_bullets": {"../../unknown": "accepted"},
            },
        )
        assert response.status_code == 422
        current = client.get(review_url).json()
        assert current["summary"]["decision"] == "pending"


def test_review_state_is_isolated_between_analysis_ids(settings):
    report = reviewable_report().model_dump(mode="json")
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        first = completed_submission(client)
        second = completed_submission(client)
        first_url = f"/api/analyses/{first['analysis_id']}/resume-review"
        second_url = f"/api/analyses/{second['analysis_id']}/resume-review"
        assert client.patch(first_url, json={"summary": "accepted"}).status_code == 200
        assert client.get(first_url).json()["summary"]["decision"] == "accepted"
        assert client.get(second_url).json()["summary"]["decision"] == "pending"


def test_final_resume_api_and_page_reflect_current_decisions(settings):
    report = reviewable_report().model_dump(mode="json")
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        payload = completed_submission(client)
        review_url = f"/api/analyses/{payload['analysis_id']}/resume-review"
        review = client.get(review_url).json()
        accepted_bullet = review["experience_bullets"][0]["id"]
        accepted_achievement = review["experience_bullets"][2]["id"]
        client.patch(
            review_url,
            json={
                "summary": "accepted",
                "experience_bullets": {
                    accepted_bullet: "accepted",
                    accepted_achievement: "rejected",
                },
            },
        )
        final_url = f"/api/analyses/{payload['analysis_id']}/final-resume"
        final = client.get(final_url)
        assert final.status_code == 200
        assert final.json()["summary"] == "Python engineer delivering reliable APIs."
        assert final.json()["experience"][0]["responsibilities"][0] == (
            "Built reliable Python APIs."
        )
        assert final.json()["experience"][0]["achievements"][0] == ("Reduced response times.")

        page = client.get(f"/results/{payload['analysis_id']}/final-resume")
        assert page.status_code == 200
        assert "Final Resume Preview" in page.text
        assert "Python engineer delivering reliable APIs." in page.text
        assert "Reduced response times." in page.text
        assert "Reduced API response times." not in page.text
        assert "← Back to Review Changes" in page.text
        assert "Download Resume" in page.text
        assert 'id="template-modal"' in page.text
        assert 'id="generate-resume" disabled' in page.text
        assert "ATS compatibility" not in page.text
        assert "Parsing integrity" not in page.text
        assert "Ollama" not in page.text


def test_review_controls_render_only_for_valid_proposals(settings):
    report = reviewable_report().model_dump(mode="json")
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        payload = completed_submission(client)
        page = client.get(payload["page_url"])
    assert page.status_code == 200
    analysis_id = payload["analysis_id"]
    assert f'data-results-page data-analysis-id="{analysis_id}"' in page.text
    assert f'data-review-url="/api/analyses/{analysis_id}/resume-review"' in page.text
    assert "/static/js/results.js?v=2.1.0" in page.text
    assert 'data-review-kind="summary"' in page.text
    assert 'data-review-kind="bullet"' in page.text
    assert 'data-review-kind="skills"' in page.text
    assert "Review Final Resume" in page.text
    assert "Review / Export Resume" in page.text
    assert f'href="/results/{analysis_id}/final-resume">Review / Export Resume</a>' in page.text
    assert "Accept" in page.text and "Reject" in page.text
    actionable = re.findall(r'<button[^>]+class="[^"]*review-decision[^"]*"[^>]*>', page.text)
    review = FinalResumeBuilder(reviewable_report()).review_payload().model_dump(mode="json")
    valid_ids = {
        "summary",
        "skills",
        *(item["id"] for item in review["experience_bullets"] if item["can_accept"]),
    }
    assert len(actionable) == len(valid_ids) * 2
    assert all('type="button"' in tag for tag in actionable)
    assert all(
        'data-decision="accepted"' in tag or 'data-decision="rejected"' in tag for tag in actionable
    )
    for item_id in valid_ids:
        assert page.text.count(f'data-review-id="{item_id}"') == 2
    invalid_id = next(item["id"] for item in review["experience_bullets"] if not item["can_accept"])
    assert f'data-review-item="{invalid_id}"' in page.text
    assert f'data-review-id="{invalid_id}"' not in page.text


def test_results_javascript_matches_rendered_review_contract():
    source = Path("resume_analyzer/web/static/js/results.js").read_text(encoding="utf-8")
    assert 'document.querySelector("[data-results-page]")' in source
    assert "window.bootstrap?.Tab" in source
    assert 'page.querySelector("[data-review-url]")' in source
    assert 'reviewSection?.addEventListener("click"' in source
    assert 'event.target.closest?.(".review-decision")' in source
    assert 'button.closest("[data-review-item]")' in source
    assert 'method: "PATCH"' in source
    assert '"Content-Type": "application/json"' in source
    assert "if (!response.ok)" in source
    assert "response.status !== 204" in source
    assert "The decision could not be saved. Please try again." in source


def test_review_accept_reject_refresh_final_page_and_pending_resolution(settings):
    report = reviewable_report().model_dump(mode="json")
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(report))) as client:
        payload = completed_submission(client)
        analysis_id = payload["analysis_id"]
        review_url = f"/api/analyses/{analysis_id}/resume-review"
        initial = client.get(review_url).json()
        accepted_id = initial["experience_bullets"][0]["id"]
        rejected_id = initial["experience_bullets"][2]["id"]

        accepted = client.patch(
            review_url,
            json={"summary": "accepted", "experience_bullets": {accepted_id: "accepted"}},
        )
        assert accepted.status_code == 200
        refreshed = client.get(review_url).json()
        assert refreshed["summary"]["decision"] == "accepted"
        assert refreshed["experience_bullets"][0]["decision"] == "accepted"

        rejected = client.patch(
            review_url,
            json={"experience_bullets": {rejected_id: "rejected"}},
        )
        assert rejected.status_code == 200
        refreshed = client.get(review_url).json()
        assert refreshed["experience_bullets"][2]["decision"] == "rejected"
        assert refreshed["skills"]["decision"] == "pending"

        refreshed_page = client.get(f"/results/{analysis_id}")
        assert refreshed_page.status_code == 200
        assert 'status-accepted" data-review-status="summary">Accepted' in refreshed_page.text
        assert (
            f'status-rejected" data-review-status="{rejected_id}">Rejected' in refreshed_page.text
        )
        assert 'class="review-choice is-final-choice" data-review-content="proposed"' in (
            refreshed_page.text
        )
        assert 'status-pending" data-review-status="skills">Pending' in refreshed_page.text

        final = client.get(f"/api/analyses/{analysis_id}/final-resume")
        assert final.status_code == 200
        final_payload = final.json()
        assert final_payload["summary"] == "Python engineer delivering reliable APIs."
        assert final_payload["experience"][0]["responsibilities"][0] == (
            "Built reliable Python APIs."
        )
        assert final_payload["experience"][0]["achievements"][0] == "Reduced response times."
        assert [group["group"] for group in final_payload["skills"]] == ["Skills"]

        page = client.get(f"/results/{analysis_id}/final-resume")
        assert page.status_code == 200
        assert f'data-final-resume-page data-analysis-id="{analysis_id}"' in page.text
        assert "/static/js/final_resume.js?v=2.1.0" in page.text
        assert "Python engineer delivering reliable APIs." in page.text
        assert "Built reliable Python APIs." in page.text
        assert "Reduced response times." in page.text
        assert "Reduced API response times." not in page.text
        assert "Programming" not in page.text
        assert "Databases" not in page.text


def test_template_api_exposes_only_allowlisted_metadata_and_previews(settings, rich_report):
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(rich_report))) as client:
        response = client.get("/api/resume-templates")
        assert response.status_code == 200
        assert [item["id"] for item in response.json()] == ["template-1", "template-2"]
        assert "docx_path" not in response.text and "C:\\Users" not in response.text
        for item in response.json():
            preview = client.get(item["preview_url"])
            assert preview.status_code == 200
            assert preview.headers["content-type"].startswith("image/")
        assert client.get("/api/resume-templates/../../private/preview").status_code == 404


@pytest.mark.parametrize("template_id", ["template-1", "template-2"])
def test_docx_download_uses_final_resume_without_rerunning_pipeline(settings, template_id):
    report = reviewable_report().model_dump(mode="json")
    factory = RecordingFactory(report)
    with TestClient(create_app(settings, pipeline_factory=factory)) as client:
        payload = completed_submission(client)
        review_url = f"/api/analyses/{payload['analysis_id']}/resume-review"
        review = client.get(review_url).json()
        client.patch(
            review_url,
            json={
                "summary": "accepted",
                "experience_bullets": {
                    review["experience_bullets"][0]["id"]: "accepted",
                    review["experience_bullets"][2]["id"]: "rejected",
                },
            },
        )
        calls_before_export = factory.calls
        response = client.post(
            f"/api/analyses/{payload['analysis_id']}/download-docx",
            json={"template_id": template_id},
        )
        assert response.status_code == 200
        assert response.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert "Jane-Doe-Resume.docx" in response.headers["content-disposition"]
        text = "\n".join(
            paragraph.text for paragraph in Document(io.BytesIO(response.content)).paragraphs
        )
        document = Document(io.BytesIO(response.content))
        text += "\n" + "\n".join(
            paragraph.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        )
        assert "Python engineer delivering reliable APIs." in text
        assert "Built reliable Python APIs." in text
        assert "Reduced response times." in text
        assert "Reduced API response times." not in text
        assert factory.calls == calls_before_export


def test_unknown_template_and_unavailable_analysis_fail_safely(settings, rich_report):
    with TestClient(create_app(settings, pipeline_factory=RecordingFactory(rich_report))) as client:
        payload = completed_submission(client)
        response = client.post(
            f"/api/analyses/{payload['analysis_id']}/download-docx",
            json={"template_id": "../../../private.docx"},
        )
        assert response.status_code == 422
        assert "C:\\Users" not in response.text
        missing = client.post(
            "/api/analyses/00000000-0000-4000-8000-000000000000/download-docx",
            json={"template_id": "template-1"},
        )
        assert missing.status_code == 404
