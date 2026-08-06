from __future__ import annotations

import contextlib
import copy
import io
import json
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches, RGBColor
from PIL import Image, ImageDraw

from resume_analyzer.extraction.contact import ContactResolver
from resume_analyzer.extraction.docx_structure_analyzer import analyze_docx_package
from resume_analyzer.extraction.duplicate_content_cleaner import deduplicate_blocks
from resume_analyzer.extraction.languages_extractor import LanguagesExtractor
from resume_analyzer.extraction.reporting_policy import print_resume_pipeline_report
from resume_analyzer.extraction.result_quality_refiner import refine_resume_result
from resume_analyzer.extraction.section_extractor import SectionExtractor
from resume_analyzer.extraction.skills_extractor import SkillsExtractor
from resume_analyzer.extraction.text_extractor import (
    TextExtractor,
    _meaningful_text_box_overlap,
)

ROOT = Path(__file__).resolve().parent
FIXTURES = ROOT / "fixtures"


def build_visual_docx(path: Path) -> Path:
    photo_path = path.with_name("candidate_photo.png")
    icon_path = path.with_name("contact_icon.png")

    photo = Image.new("RGB", (300, 400), "white")
    photo_draw = ImageDraw.Draw(photo)
    photo_draw.rectangle(
        (50, 40, 250, 360),
        fill=(200, 160, 140),
    )
    photo.save(photo_path)

    icon = Image.new("RGBA", (32, 32), (0, 0, 0, 0))
    icon_draw = ImageDraw.Draw(icon)
    icon_draw.ellipse(
        (4, 4, 28, 28),
        fill=(0, 114, 188, 255),
    )
    icon.save(icon_path)

    document = Document()

    mirrored_values = (
        ["JORDAN EXAMPLE"] * 2
        + ["OBJECTIVE"] * 2
        + [
            "Describe in a few lines your career path, "
            "your key skills for the position and your career goals."
        ]
        * 2
        + ["SKILLS"] * 2
        + ["Project management Strong decision maker " "Complex problem solver Creative design"] * 2
        + ["EXPERIENCE"] * 2
        + ["EDUCATION"] * 2
        + ["Bachelor of Science: Computer Information Systems | " "2014 Columbia University, NY"]
        * 2
        + ["LANGUAGES"] * 2
        + ["Spanish - C2 Chinese - A1 German - A2"] * 2
        + ["HOBBIES"] * 2
        + ["Writing Sketching Photography"] * 2
    )

    for value in mirrored_values:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(value)
        if value in {
            "JORDAN EXAMPLE",
            "SKILLS",
            "EXPERIENCE",
        }:
            run.font.color.rgb = RGBColor(0, 114, 188)

    placeholder = (
        "Job Title Company Name, Location | Jan 2020 - current "
        "Key responsibility or achievement "
        "Key responsibility or achievement "
        "Key responsibility or achievement "
        "Key responsibility or achievement"
    )
    for _ in range(6):
        document.add_paragraph(placeholder)

    for value in (
        "GRAPHIC DESIGNER",
        "London, England",
        "+1 555 010 0300",
        "jordan.visual @example.test",
    ):
        document.add_paragraph(value)

    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(
        str(photo_path),
        width=Inches(1.2),
    )
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(
        str(icon_path),
        width=Inches(0.15),
    )

    textboxes = (
        ("OBJECTIVE", 30, 20, "#0072BC"),
        ("SKILLS", 30, 120, "#00A14B"),
        ("EXPERIENCE", 300, 120, "#E82C2A"),
        ("CONTACT", 300, 20, "#FDB913"),
    )
    for index, (value, left, top, color) in enumerate(
        textboxes,
        start=1,
    ):
        xml = f"""<w:r
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml">
          <w:pict>
            <v:shape id="tb{index}"
              style="position:absolute;margin-left:{left}pt;
                     margin-top:{top}pt;width:180pt;height:40pt"
              fillcolor="{color}">
              <v:textbox>
                <w:txbxContent>
                  <w:p><w:r><w:t>{value}</w:t></w:r></w:p>
                </w:txbxContent>
              </v:textbox>
            </v:shape>
          </w:pict>
        </w:r>"""
        document.add_paragraph()._p.append(parse_xml(xml))

    document.save(path)
    return path


def load_visual_output() -> dict:
    return json.loads((FIXTURES / "visual_docx_current_output.json").read_text(encoding="utf-8"))


def test_docx_ooxml_assets_colors_and_photo(tmp_path: Path) -> None:
    path = build_visual_docx(tmp_path / "visual.docx")
    analysis = analyze_docx_package(path)

    assert analysis["status"] == "ok"
    style = analysis["document_style"]
    assert style["has_color"] is True
    assert style["chromatic_color_count"] >= 4
    assert style["fixed_palette_used"] is False
    assert style["palette_method"] == ("docx_ooxml_theme_direct_and_shape_analysis")

    assets = analysis["document_assets"]
    assert assets["image_count"] == 2
    assert assets["icon_count"] >= 1
    assert assets["candidate_photo_detected"] is True
    assert assets["text_box_count"] == 4


def test_docx_duplicate_cleanup_preserves_logical_slots(
    tmp_path: Path,
) -> None:
    path = build_visual_docx(tmp_path / "duplicates.docx")
    analysis = analyze_docx_package(path)
    cleaned = deduplicate_blocks(analysis["visual_blocks"])
    duplicate = cleaned["duplicate_analysis"]

    assert duplicate["mirror_factor"] == 2
    assert duplicate["duplicate_ratio"] >= 0.40
    assert duplicate["logical_placeholder_slots"][0]["logical_slot_count"] == 3
    assert sum("Job Title Company Name" in item["text"] for item in cleaned["blocks"]) == 3

    extraction = TextExtractor().extract(str(path))
    assert extraction["success"] is True
    assert extraction["engine"] == "docx_ooxml"
    assert sum("Job Title Company Name" in line for line in extraction["text"].splitlines()) == 3
    assert any(
        value.startswith("removed_docx_duplicate_blocks:") for value in extraction["warnings"]
    )


def test_overlap_threshold_rejects_font_box_bleed_but_keeps_real_overlay() -> None:
    first = (40.0, 67.1, 128.0, 83.6)
    normally_spaced_next_line = (40.0, 80.3, 145.4, 96.8)
    materially_overlaid_line = (40.0, 73.1, 145.4, 89.6)

    assert not _meaningful_text_box_overlap(first, normally_spaced_next_line)
    assert _meaningful_text_box_overlap(first, materially_overlaid_line)


def test_accounting_pdf_does_not_treat_normal_line_spacing_as_overlap() -> None:
    extraction = TextExtractor(enable_ocr=False).extract(str(FIXTURES / "resume-accounting.pdf"))
    visual = extraction["visual_metadata"]

    assert extraction["success"] is True
    assert visual["overlap_count"] == 0
    assert visual["repeated_header_footer_count"] == 2
    assert "removed_repeated_header_footer_blocks:2" in extraction["warnings"]


def test_visual_docx_semantic_reconciliation() -> None:
    result = refine_resume_result(
        copy.deepcopy(load_visual_output()),
        copy_result=False,
    )

    profile = result["document_profile"]
    assert profile["document_type"] == ("partially_completed_resume_template")
    assert profile["placeholder_role_slot_count"] == 3
    assert set(profile["unresolved_template_sections"]) == {
        "objective",
        "experience",
    }

    duplicate = result["duplicate_analysis"]
    assert duplicate["raw_block_count"] == 44
    assert duplicate["logical_placeholder_slots"][0]["logical_slot_count"] == 3

    contact = result["contact"]
    assert contact["email"] == "jordan.visual@example.test"
    assert contact["email_raw"] == "jordan.visual @example.test"
    assert contact["email_normalization"] == ("removed_internal_style_whitespace")
    assert contact["phone"] == "+15550100300"
    assert contact["job_title"] == "Graphic Designer"

    assert result["skills"]["all_skills"] == [
        "Project Management",
        "Creative Design",
        "Complex Problem Solving",
        "Decision Making",
    ]
    assert "NY" not in result["skills"]["all_skills"]
    assert "Computer Information Systems" not in (result["skills"]["all_skills"])

    languages = {
        item["language"]: (
            item["cefr"],
            item["proficiency"],
        )
        for item in result["languages"]["languages"]
    }
    assert languages == {
        "Spanish": ("C2", "Proficient"),
        "Chinese": ("A1", "Beginner"),
        "German": ("A2", "Elementary"),
    }

    experience = result["experience"]
    assert experience["count"] == 0
    assert experience["placeholder_role_slot_count"] == 3
    assert experience["professional_duration_status"] == ("not_computable_template_placeholders")

    assert result["source_readiness"]["status"] == ("template_incomplete")
    assert result["source_readiness"]["trusted"] is False
    assert result["summary"]["languages_count"] == 3
    assert result["summary"]["job_title"] == ("Graphic Designer")


def test_pairwise_language_extractor_does_not_share_first_level() -> None:
    payload = {
        "cleaned_text": ("LANGUAGES\n" "Spanish - C2 Chinese - A1 German - A2"),
        "sections": {"languages": {"content": ("Spanish - C2 Chinese - A1 German - A2")}},
    }
    result = LanguagesExtractor().extract(payload)
    assert [item["cefr"] for item in result["languages"]] == [
        "C2",
        "A1",
        "A2",
    ]


def test_contact_normalizes_styled_email_and_international_phone() -> None:
    text = """JORDAN EXAMPLE
GRAPHIC DESIGNER
London, England
+1 555 010 0300
jordan.visual @example.test
"""
    contact = ContactResolver().resolve(
        text=text,
        raw_text=text,
        layout_blocks=[],
    )
    assert contact["email"] == "jordan.visual@example.test"
    assert contact["phone"] == "+15550100300"
    assert contact["job_title"] == "Graphic Designer"


def test_section_repair_moves_semantic_leakage() -> None:
    text = """JORDAN EXAMPLE
HOBBIES
Writing Sketching Photography
Job Title Company Name, Location | Jan 2020 - current Key responsibility or achievement
EDUCATION
Bachelor of Science: Information Systems | 2014 University
Spanish - C2 Chinese - A1 German - A2
LANGUAGES
Writing Sketching Photography
EXPERIENCE
GRAPHIC DESIGNER
London, England
jordan.visual @example.test
Describe in a few lines your career path and your career goals.
OBJECTIVE
"""
    result = SectionExtractor().extract_sections(text)
    sections = result["sections"]
    assert "Job Title Company Name" in sections["experience"]["content"]
    assert "Spanish - C2" in sections["languages"]["content"]
    assert "Writing Sketching Photography" in (sections["interests"]["content"])
    assert "Describe in a few lines" in (sections["summary"]["content"])


def test_skills_compound_visual_line() -> None:
    payload = {
        "sections": {
            "skills": {
                "content": (
                    "Project management Strong decision maker "
                    "Complex problem solver Creative design"
                )
            },
            "education": {
                "content": (
                    "Bachelor of Science: Computer Information " "Systems | Columbia University, NY"
                )
            },
            "summary": {"content": ""},
            "experience": {"content": ""},
            "projects": {"content": ""},
        }
    }
    result = SkillsExtractor(
        use_spacy=False,
        use_sbert=False,
    ).extract(payload)
    assert {
        "Project Management",
        "Decision Making",
        "Complex Problem Solving",
        "Creative Design",
    } <= set(result["all_skills"])
    assert "NY" not in result["all_skills"]


def test_visual_ats_reporting_section() -> None:
    result = refine_resume_result(
        copy.deepcopy(load_visual_output()),
        copy_result=False,
    )
    result["document_style"] = {
        "has_color": True,
        "chromatic_color_count": 6,
    }
    result["document_assets"] = {
        "image_count": 2,
        "candidate_photo_detected": True,
        "icon_count": 3,
        "text_box_count": 8,
    }
    result["ats_structure"] = {
        "risk_level": "high",
    }

    stream = io.StringIO()
    with contextlib.redirect_stdout(stream):
        print_resume_pipeline_report(result)
    output = stream.getvalue()

    assert "🎨 Visual / ATS Structure:" in output
    assert "Colors Detected:       yes (6 chromatic)" in output
    assert "Embedded Images:       2" in output
    assert "Candidate Photo:       detected" in output
    assert "Template Status:       partially_completed_resume_template" in output
    assert "ATS Visual Risk:       high" in output
    assert "Professional Duration: not computable — template role placeholders" in output


def test_new_source_files_have_no_resume_specific_identity_hardcoding() -> None:
    source_root = ROOT.parent
    extraction_root = source_root / "resume_analyzer" / "extraction"
    files = [
        extraction_root / "docx_structure_analyzer.py",
        extraction_root / "duplicate_content_cleaner.py",
        extraction_root / "text_extractor.py",
        extraction_root / "document_intelligence.py",
        extraction_root / "section_extractor.py",
        extraction_root / "languages_extractor.py",
    ]
    source = "\n".join(path.read_text(encoding="utf-8") for path in files).casefold()
    for forbidden in (
        "jordan example",
        "free-english-cv-for-women",
        "columbia university",
        "jordan example",
        "jordan example",
        "mohammed ali",
    ):
        assert forbidden not in source
