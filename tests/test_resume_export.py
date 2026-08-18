from __future__ import annotations

from copy import deepcopy
from hashlib import sha256
from io import BytesIO

import pytest
from docx import Document
from pydantic import ValidationError

from resume_analyzer.export import (
    DEFAULT_TEMPLATE_REGISTRY,
    DocxRenderer,
    FinalCertification,
    FinalContact,
    FinalEducation,
    FinalExperience,
    FinalLanguage,
    FinalProject,
    FinalResume,
    FinalResumeBuilder,
    FinalSkillGroup,
    ReviewDecision,
    ReviewStateError,
    ReviewUpdate,
    TemplateNotFound,
    content_disposition,
)
from resume_analyzer.export.final_resume_builder import bullet_review_id
from resume_analyzer.schemas import PipelineReport
from tests.report_fixtures import make_report


def reviewable_report() -> PipelineReport:
    report = make_report(
        experience=[
            {
                "job_title": "Software Engineer",
                "company": "Acme Corp",
                "location": "Example City",
                "start_date": "2021",
                "end_date": "Present",
                "current": True,
                "responsibilities": [
                    "Built APIs.",
                    "Maintained services.",
                    "Supported releases.",
                ],
                "achievements": ["Reduced response times."],
                "technologies": ["Python", "SQL"],
                "confidence": 0.95,
            }
        ],
        languages=[{"language": "English", "proficiency": "Fluent"}],
        certifications=[{"name": "Example Certificate", "issuer": "Example Org"}],
    )
    payload = report.model_dump(mode="json")
    payload["rewrites"] = {
        "status": "complete",
        "language": "en",
        "provider": "mock",
        "model": "synthetic",
        "summary": {
            "status": "improved",
            "original": payload["entities"]["summary"],
            "improved": "Python engineer delivering reliable APIs.",
        },
        "experience_bullets": [
            {
                "experience_index": 0,
                "bullet_index": 0,
                "bullet_kind": "responsibility",
                "status": "improved",
                "original": "Built APIs.",
                "improved": "Built reliable Python APIs.",
            },
            {
                "experience_index": 0,
                "bullet_index": 1,
                "bullet_kind": "responsibility",
                "status": "rejected",
                "original": "Maintained services.",
                "improved": None,
                "warnings": ["Unsupported factual claim."],
            },
            {
                "experience_index": 0,
                "bullet_index": 0,
                "bullet_kind": "achievement",
                "status": "improved",
                "original": "Reduced response times.",
                "improved": "Reduced API response times.",
            },
        ],
        "skills_section": {
            "status": "improved",
            "method": "ai",
            "original_items": ["Python", "SQL", "Postgres"],
            "improved_groups": [
                {"group": "Programming", "items": ["Python"]},
                {"group": "Databases", "items": ["SQL", "Postgres"]},
            ],
        },
        "completed_components": ["summary", "experience_bullets", "skills_section"],
        "bullet_stats": {"total_eligible": 4, "selected": 3, "processed": 3, "skipped": 1},
    }
    payload["module_status"]["rewrites"] = {
        "status": "complete",
        "provider": "mock",
        "model": "synthetic",
        "detail": None,
    }
    return PipelineReport.model_validate(payload)


def test_pending_review_defaults_every_proposal_to_original() -> None:
    builder = FinalResumeBuilder(reviewable_report())
    final = builder.build(builder.initial_state())
    assert final.summary == builder.report.entities.summary
    assert final.experience[0].responsibilities == [
        "Built APIs.",
        "Maintained services.",
        "Supported releases.",
    ]
    assert final.experience[0].achievements == ["Reduced response times."]
    assert [item for group in final.skills for item in group.items] == [
        "Python",
        "SQL",
        "Postgres",
    ]


@pytest.mark.parametrize(
    ("decision", "expected"),
    [
        (ReviewDecision.ACCEPTED, "Python engineer delivering reliable APIs."),
        (ReviewDecision.REJECTED, "Python developer building APIs."),
        (ReviewDecision.PENDING, "Python developer building APIs."),
    ],
)
def test_summary_decision_resolution(decision: ReviewDecision, expected: str) -> None:
    builder = FinalResumeBuilder(reviewable_report())
    state = builder.apply_update(builder.initial_state(), ReviewUpdate(summary=decision))
    assert builder.build(state).summary == expected


def test_accepted_bullet_is_independent_and_order_is_preserved() -> None:
    builder = FinalResumeBuilder(reviewable_report())
    item_id = bullet_review_id(0, "responsibility", 0)
    state = builder.apply_update(
        builder.initial_state(),
        ReviewUpdate(experience_bullets={item_id: ReviewDecision.ACCEPTED}),
    )
    assert builder.build(state).experience[0].responsibilities == [
        "Built reliable Python APIs.",
        "Maintained services.",
        "Supported releases.",
    ]


def test_rejected_or_unprocessed_bullets_cannot_be_accepted() -> None:
    builder = FinalResumeBuilder(reviewable_report())
    with pytest.raises(ReviewStateError, match="Unknown or unavailable"):
        builder.apply_update(
            builder.initial_state(),
            ReviewUpdate(
                experience_bullets={
                    bullet_review_id(0, "responsibility", 1): ReviewDecision.ACCEPTED
                }
            ),
        )
    with pytest.raises(ReviewStateError, match="Unknown or unavailable"):
        builder.apply_update(
            builder.initial_state(),
            ReviewUpdate(
                experience_bullets={
                    bullet_review_id(0, "responsibility", 2): ReviewDecision.ACCEPTED
                }
            ),
        )


@pytest.mark.parametrize(
    ("decision", "expected_groups"),
    [
        (ReviewDecision.ACCEPTED, ["Programming", "Databases"]),
        (ReviewDecision.REJECTED, ["Skills"]),
        (ReviewDecision.PENDING, ["Skills"]),
    ],
)
def test_skills_decision_resolution(decision: ReviewDecision, expected_groups: list[str]) -> None:
    builder = FinalResumeBuilder(reviewable_report())
    state = builder.apply_update(builder.initial_state(), ReviewUpdate(skills=decision))
    assert [group.group for group in builder.build(state).skills] == expected_groups


def test_unsupported_or_omitted_skills_disable_the_proposal() -> None:
    payload = reviewable_report().model_dump(mode="json")
    payload["rewrites"]["skills_section"]["improved_groups"] = [
        {"group": "Cloud", "items": ["AWS"]}
    ]
    builder = FinalResumeBuilder(PipelineReport.model_validate(payload))
    assert builder.review_payload().skills.can_accept is False
    with pytest.raises(ReviewStateError, match="no acceptable proposal"):
        builder.apply_update(builder.initial_state(), ReviewUpdate(skills=ReviewDecision.ACCEPTED))


def test_invalid_and_no_material_change_summary_remain_original() -> None:
    for status in ("rejected", "unchanged", "unavailable", "failed" if False else "not_run"):
        payload = reviewable_report().model_dump(mode="json")
        payload["rewrites"]["summary"].update(status=status, improved=None)
        if status == "not_run":
            payload["rewrites"]["summary"]["original"] = ""
        builder = FinalResumeBuilder(PipelineReport.model_validate(payload))
        review = builder.review_payload()
        assert review.summary.can_accept is False
        assert review.summary.final == payload["entities"]["summary"]


def test_final_resume_preserves_all_untouched_semantic_entities() -> None:
    final = FinalResumeBuilder(reviewable_report()).build()
    assert final.contact.name == "Jane Doe"
    assert final.contact.email == "jane@example.com"
    assert final.experience[0].job_title == "Software Engineer"
    assert final.education[0].institution == "Example University"
    assert final.projects[0].name == "API Project"
    assert final.languages[0].language == "English"
    assert final.certifications[0].name == "Example Certificate"
    assert FinalResume.model_validate(final.model_dump(mode="json")) == final


def test_original_report_is_not_mutated_by_review_or_build() -> None:
    report = reviewable_report()
    before = deepcopy(report.model_dump(mode="json"))
    builder = FinalResumeBuilder(report)
    state = builder.apply_update(
        builder.initial_state(), ReviewUpdate(summary=ReviewDecision.ACCEPTED)
    )
    builder.build(state)
    assert report.model_dump(mode="json") == before


def test_review_update_rejects_unknown_fields_and_decisions() -> None:
    with pytest.raises(ValidationError):
        ReviewUpdate.model_validate({"summary": "approved"})
    with pytest.raises(ValidationError):
        ReviewUpdate.model_validate({"unknown": "accepted"})


def test_template_registry_is_allowlisted_and_public_metadata_has_no_paths() -> None:
    metadata = DEFAULT_TEMPLATE_REGISTRY.public_metadata()
    assert [item["id"] for item in metadata] == ["template-1", "template-2"]
    assert all("docx_path" not in item and "preview_path" not in item for item in metadata)
    for definition in DEFAULT_TEMPLATE_REGISTRY.definitions():
        assert definition.docx_path.is_file()
        assert definition.preview_path.is_file()
    with pytest.raises(TemplateNotFound):
        DEFAULT_TEMPLATE_REGISTRY.get("../../private.docx")


def _docx_text(value: bytes) -> str:
    document = Document(BytesIO(value))
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    return "\n".join(values)


def _complete_resume() -> FinalResume:
    return FinalResume(
        contact=FinalContact(
            name="Alex Example",
            job_title="Principal Engineer",
            email="alex@example.test",
            phone="+1 555 010 0000",
            location="Example City",
            linkedin="https://linkedin.example/alex",
            github="https://github.example/alex",
            portfolio="https://alex.example",
        ),
        summary="Final accepted summary for the synthetic candidate.",
        experience=[
            FinalExperience(
                job_title="Principal Engineer",
                company="Example Company",
                location="Example City",
                start_date="2022",
                end_date="Present",
                current=True,
                responsibilities=["Accepted final bullet.", "Original rejected bullet."],
                achievements=["Improved a documented process."],
                technologies=["Python", "FastAPI"],
                metrics=["Synthetic metric label"],
            )
        ],
        education=[
            FinalEducation(
                degree="BSc Computer Science",
                institution="Example University",
                graduation_year=2021,
                gpa="3.8",
                honors=["Example Honor"],
                coursework=["Distributed Systems"],
            )
        ],
        skills=[
            FinalSkillGroup(group="Programming", items=["Python", "SQL"]),
            FinalSkillGroup(group="Tools", items=["Docker"]),
        ],
        projects=[
            FinalProject(
                name="Synthetic Project",
                role="Lead Developer",
                description="Built a deterministic test application.",
                technologies=["Python"],
                url="https://project.example",
            )
        ],
        languages=[FinalLanguage(language="English", proficiency="Fluent", cefr="C2")],
        certifications=[
            FinalCertification(
                name="Synthetic Certification",
                issuer="Example Institute",
                date="2025",
                credential_id="EX-123",
                url="https://certificate.example",
            )
        ],
    )


@pytest.mark.parametrize("template_id", ["template-1", "template-2"])
def test_docx_renderer_uses_final_content_and_preserves_template_bytes(template_id: str) -> None:
    renderer = DocxRenderer()
    definition = DEFAULT_TEMPLATE_REGISTRY.get(template_id)
    before = sha256(definition.docx_path.read_bytes()).digest()
    value = renderer.render(_complete_resume(), template_id)
    text = _docx_text(value)
    Document(BytesIO(value))
    assert "Alex Example" in text
    assert "Final accepted summary for the synthetic candidate." in text
    assert "Accepted final bullet." in text
    assert "Original rejected bullet." in text
    assert "Rejected proposed bullet." not in text
    assert "Synthetic Project" in text
    assert "Example University" in text
    assert "Synthetic Certification" in text
    assert "{{" not in text and "{%" not in text
    assert sha256(definition.docx_path.read_bytes()).digest() == before


@pytest.mark.parametrize("template_id", ["template-1", "template-2"])
def test_short_resume_omits_empty_optional_sections(template_id: str) -> None:
    resume = FinalResume(
        contact=FinalContact(
            name="Minimal Example",
            email="minimal@example.test",
            phone="+1 555 010 0010",
        ),
        summary="A concise synthetic summary.",
        education=[FinalEducation(degree="BSc", institution="Example University", end_date="2024")],
        skills=[FinalSkillGroup(group="Skills", items=["Python", "SQL"])],
    )
    text = _docx_text(DocxRenderer().render(resume, template_id))
    assert "Minimal Example" in text
    assert "PROJECTS" not in text.upper()
    assert "CERTIFICATIONS" not in text.upper()
    assert "LANGUAGES" not in text.upper()
    assert "None" not in text and "null" not in text


def _long_resume() -> FinalResume:
    resume = _complete_resume()
    resume.summary = " ".join(["Long synthetic summary content."] * 35)
    resume.experience = [
        FinalExperience(
            job_title=f"Synthetic Role {index}",
            company=f"Example Company {index}",
            start_date=str(2015 + index),
            end_date=str(2016 + index),
            responsibilities=[
                f"Experience {index} final bullet {bullet}. "
                "Delivered a fully synthetic, deterministic result."
                for bullet in range(5)
            ],
            technologies=["Python", "SQL", f"Tool-{index}"],
        )
        for index in range(5)
    ]
    resume.projects = [
        FinalProject(
            name=f"Synthetic Project {index}",
            description=f"Project description marker {index} with deterministic content.",
            technologies=["Python", f"Library-{index}"],
        )
        for index in range(8)
    ]
    resume.skills = [
        FinalSkillGroup(group="Synthetic Skills", items=[f"Skill-{index}" for index in range(40)])
    ]
    resume.education = [
        FinalEducation(degree=f"Synthetic Degree {index}", institution="Example University")
        for index in range(2)
    ]
    resume.certifications = [
        FinalCertification(name=f"Synthetic Certification {index}") for index in range(3)
    ]
    resume.languages = [
        FinalLanguage(language=value, proficiency="Professional")
        for value in ("English", "Arabic", "French")
    ]
    return resume


@pytest.mark.parametrize("template_id", ["template-1", "template-2"])
def test_long_resume_preserves_every_repeated_item(template_id: str) -> None:
    text = _docx_text(DocxRenderer().render(_long_resume(), template_id))
    for index in range(5):
        assert f"Synthetic Role {index}" in text
        for bullet in range(5):
            assert f"Experience {index} final bullet {bullet}." in text
    for index in range(8):
        assert f"Synthetic Project {index}" in text
    for index in range(40):
        assert f"Skill-{index}" in text
    for index in range(3):
        assert f"Synthetic Certification {index}" in text


@pytest.mark.parametrize("template_id", ["template-1", "template-2"])
def test_arabic_unicode_survives_docx_rendering(template_id: str) -> None:
    resume = FinalResume(
        contact=FinalContact(name="ليلى أحمد", email="layla@example.test"),
        summary="مهندسة برمجيات تطور تطبيقات موثوقة.",
        experience=[
            FinalExperience(
                job_title="مهندسة برمجيات",
                company="شركة مثال",
                responsibilities=["طورت خدمات آمنة باستخدام بايثون."],
            )
        ],
        skills=[FinalSkillGroup(group="المهارات", items=["بايثون", "تحليل البيانات"])],
    )
    text = _docx_text(DocxRenderer().render(resume, template_id))
    assert "ليلى أحمد" in text
    assert "مهندسة برمجيات تطور تطبيقات موثوقة." in text
    assert "طورت خدمات آمنة باستخدام بايثون." in text
    assert "تحليل البيانات" in text


def test_content_disposition_supports_arabic_and_blocks_header_injection() -> None:
    header = content_disposition("ليلى أحمد\r\nX-Evil: injected/..\\name")
    assert "filename*=UTF-8''" in header
    assert "%D9%84" in header
    assert "\r" not in header and "\n" not in header
    assert "X-Evil:" not in header
