"""Create synthetic PDF, DOCX, and scanned-PDF files for local validation."""

from __future__ import annotations

import argparse
import io
import json
from pathlib import Path

import fitz
from docx import Document
from PIL import Image, ImageDraw, ImageFont

RESUME_LINES = (
    "Jordan Example",
    "Software Engineer",
    "Summary",
    "Software engineer building reliable Python and FastAPI services.",
    "Skills",
    "Python, FastAPI, SQL, Docker, API testing",
    "Experience",
    "Software Engineer | Example Labs | 2022 - Present",
    "Built FastAPI services and wrote SQL queries for internal applications.",
    "Education",
    "Bachelor of Computer Science | Example University | 2021",
)


def _font(size: int) -> ImageFont.FreeTypeFont:
    candidates = (
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\tahoma.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    )
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size=size)
    raise RuntimeError("No Unicode TrueType font was found for synthetic OCR documents")


def _write_text_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.insert_textbox(
        fitz.Rect(54, 54, 558, 738),
        "\n".join(RESUME_LINES),
        fontsize=11,
        lineheight=1.35,
    )
    document.save(path)
    document.close()


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_heading(RESUME_LINES[0], level=0)
    document.add_paragraph(RESUME_LINES[1])
    for heading in ("Summary", "Skills", "Experience", "Education"):
        index = RESUME_LINES.index(heading)
        document.add_heading(heading, level=1)
        end = next(
            (
                RESUME_LINES.index(next_heading)
                for next_heading in ("Skills", "Experience", "Education")
                if RESUME_LINES.index(next_heading) > index
            ),
            len(RESUME_LINES),
        )
        for line in RESUME_LINES[index + 1 : end]:
            document.add_paragraph(line)
    document.save(path)


def _write_scanned_pdf(path: Path, lines: tuple[tuple[str, str], ...]) -> None:
    image = Image.new("RGB", (2480, 3508), "white")
    draw = ImageDraw.Draw(image)
    font = _font(72)
    y = 260
    for text, direction in lines:
        if direction == "rtl":
            draw.text((2220, y), text, font=font, fill="black", anchor="ra", direction="rtl")
        else:
            draw.text((260, y), text, font=font, fill="black")
        y += 150

    encoded = io.BytesIO()
    image.save(encoded, format="JPEG", quality=90, optimize=True)
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_image(page.rect, stream=encoded.getvalue())
    document.save(path)
    document.close()


def create_documents(destination: Path) -> dict[str, Path]:
    destination.mkdir(parents=True, exist_ok=True)
    paths = {
        "pdf": destination / "synthetic_resume.pdf",
        "docx": destination / "synthetic_resume.docx",
        "english_scan": destination / "synthetic_english_scan.pdf",
        "arabic_scan": destination / "synthetic_arabic_scan.pdf",
        "mixed_scan": destination / "synthetic_mixed_scan.pdf",
        "job_description": destination / "synthetic_job_description.txt",
    }
    _write_text_pdf(paths["pdf"])
    _write_docx(paths["docx"])
    _write_scanned_pdf(
        paths["english_scan"],
        (
            ("JORDAN EXAMPLE SOFTWARE ENGINEER", "ltr"),
            ("SKILLS PYTHON FASTAPI SQL DOCKER", "ltr"),
            ("EXPERIENCE BUILDING RELIABLE BACKEND API SERVICES", "ltr"),
        ),
    )
    _write_scanned_pdf(
        paths["arabic_scan"],
        (
            ("مهندس برمجيات", "rtl"),
            ("المهارات بايثون وتطوير الويب", "rtl"),
            ("خبرة في بناء خدمات وتطبيقات موثوقة", "rtl"),
        ),
    )
    _write_scanned_pdf(
        paths["mixed_scan"],
        (
            ("JORDAN EXAMPLE SOFTWARE ENGINEER", "ltr"),
            ("مهندس برمجيات وتطوير الويب", "rtl"),
            ("SKILLS PYTHON FASTAPI SQL DOCKER", "ltr"),
        ),
    )
    paths["job_description"].write_text(
        "Software engineer role requiring Python, FastAPI, SQL, Docker, and API testing.\n",
        encoding="utf-8",
    )
    return paths


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, default=Path("runtime/live-validation"))
    args = parser.parse_args()
    paths = create_documents(args.output_dir)
    relative = {key: path.as_posix() for key, path in paths.items()}
    print(json.dumps(relative, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
