"""Generate synthetic live-validation documents outside the committed test fixtures."""

from __future__ import annotations

import argparse
import random
from pathlib import Path

import fitz
from docx import Document
from PIL import Image, ImageDraw, ImageFont


def _one_column_pdf(path: Path, lines: list[str]) -> None:
    document = fitz.open()
    page = document.new_page(width=612, height=792)
    y = 42
    headings = {
        "summary",
        "profile",
        "about me",
        "skills",
        "technical skills",
        "experience",
        "work history",
        "career history",
        "projects",
        "portfolio",
        "education",
    }
    for line in lines:
        heading = line.casefold() in headings
        page.insert_text(
            (50, y),
            line,
            fontsize=12 if heading else 9.5,
            fontname="hebo" if heading else "helv",
        )
        y += 20 if heading else 15
    document.save(path)
    document.close()


def _two_column_pdf(
    path: Path,
    *,
    name: str,
    left_heading: str,
    right_heading: str,
    split: float,
) -> None:
    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.insert_text((40, 44), name, fontsize=18, fontname="hebo")
    page.insert_text((40, 64), "Systems Engineer", fontsize=10)
    page.insert_text((40, 82), "candidate@example.test | +1 555 010 8080", fontsize=9)
    page.insert_text((40, 120), left_heading, fontsize=12, fontname="hebo")
    page.insert_text((40, 142), "Engineer focused on reliable systems.", fontsize=9)
    page.insert_text((40, 180), "TECHNICAL SKILLS", fontsize=12, fontname="hebo")
    for index, value in enumerate(["Python", "SQL", "Docker", "Communication"]):
        page.insert_text((40, 202 + index * 17), value, fontsize=9)
    page.insert_text((split + 20, 120), right_heading, fontsize=12, fontname="hebo")
    right_lines = [
        "Systems Engineer",
        "Unseen Platform Group",
        "2021 - Present",
        "- Built reliable service APIs.",
        "- Added integration test coverage.",
        "PORTFOLIO",
        "Operations Console",
        "2024",
        "Built a dashboard for operational reporting.",
        "EDUCATION",
        "Example University - Bachelor of Computer Science",
        "2017 - 2021",
    ]
    y = 142
    for line in right_lines:
        heading = line in {"PORTFOLIO", "EDUCATION"}
        page.insert_text(
            (split + 20, y),
            line,
            fontsize=11 if heading else 8.7,
            fontname="hebo" if heading else "helv",
        )
        y += 20 if heading else 15
    document.save(path)
    document.close()


def _image_pdf(path: Path, lines: list[str], *, rtl: bool = False) -> None:
    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    font = (
        ImageFont.truetype(str(font_path), 30) if font_path.exists() else ImageFont.load_default()
    )
    y = 70
    for line in lines:
        x = 1170 if rtl else 80
        try:
            draw.text(
                (x, y),
                line,
                fill="black",
                font=font,
                anchor="ra" if rtl else None,
                direction="rtl" if rtl else None,
            )
        except (TypeError, ValueError):
            draw.text((80, y), line, fill="black", font=font)
        y += 55
    stream = fitz.open()
    page = stream.new_page(width=612, height=792)
    pixmap_path = path.with_suffix(".png")
    image.save(pixmap_path)
    page.insert_image(page.rect, filename=str(pixmap_path))
    stream.save(path)
    stream.close()
    pixmap_path.unlink()


def _resume_lines(name: str, company: str, project: str, year: int) -> list[str]:
    email_name = name.casefold().replace(" ", ".")
    return [
        name,
        "Platform Engineer",
        f"{email_name}@example.test | +1 555 010 {year}",
        "Austin, TX",
        "SUMMARY",
        "Engineer building reliable services and accessible products.",
        "SKILLS",
        "Python",
        "SQL",
        "Docker",
        "EXPERIENCE",
        "Platform Engineer",
        company,
        f"{year - 4} - Present",
        "- Built reliable APIs with Python and SQL.",
        "- Improved observability through measured service telemetry.",
        "PROJECTS",
        project,
        str(year),
        "Built an operational reporting service for internal teams.",
        "EDUCATION",
        "Example University - Bachelor of Computer Science",
        f"{year - 8} - {year - 4}",
    ]


def generate(output: Path, seed: int) -> list[Path]:
    output.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []
    clean_lines = _resume_lines("Robin Mercer", "Harbor Systems", "Service Health Console", 2025)
    clean_pdf = output / "clean-professional.pdf"
    _one_column_pdf(clean_pdf, clean_lines)
    generated.append(clean_pdf)

    docx_path = output / "clean-professional.docx"
    document = Document()
    for line in clean_lines:
        document.add_paragraph(line)
    document.save(docx_path)
    generated.append(docx_path)

    scanned_en = output / "scanned-english.pdf"
    _image_pdf(scanned_en, clean_lines)
    generated.append(scanned_en)

    arabic_lines = [
        "ليان منصور",
        "مهندسة برمجيات",
        "lian@example.test | +963 944 123 456",
        "الملخص المهني",
        "مهندسة تطور أنظمة موثوقة",
        "المهارات",
        "Python SQL Docker",
        "الخبرة",
        "مهندسة برمجيات",
        "شركة تقنية",
        "2021 - الآن",
    ]
    arabic_docx = output / "clean-arabic.docx"
    document = Document()
    for line in arabic_lines:
        document.add_paragraph(line)
    document.save(arabic_docx)
    generated.append(arabic_docx)

    scanned_ar = output / "scanned-arabic.pdf"
    _image_pdf(scanned_ar, arabic_lines, rtl=True)
    generated.append(scanned_ar)

    mixed = output / "mixed-language.pdf"
    _image_pdf(
        mixed,
        [
            "Nour Haddad | نور حداد",
            "Software Engineer | مهندسة برمجيات",
            "nour@example.test | +963 944 654 321",
            "SUMMARY | الملخص",
            "Builds Python services and أنظمة موثوقة",
            "SKILLS | المهارات",
            "Python SQL Docker تواصل",
        ],
    )
    generated.append(mixed)

    generator = random.Random(seed)
    first_names = ["Emery", "Sasha", "Milan", "Ari", "Noel", "Remy", "Dara"]
    last_names = ["Rowan", "Vale", "Ibarra", "Chen", "Okafor", "Petrov", "Salim"]
    companies = ["Silver Pine Labs", "Keystone Data", "Atlas River", "Lumen Works"]
    projects = ["Queue Monitor", "Demand Planner", "Release Console", "Audit Service"]
    for index in range(5):
        path = output / f"randomized-{index + 1}.pdf"
        _one_column_pdf(
            path,
            _resume_lines(
                f"{generator.choice(first_names)} {generator.choice(last_names)}",
                generator.choice(companies),
                generator.choice(projects),
                2020 + index,
            ),
        )
        generated.append(path)

    aliases = [
        ("ABOUT ME", "CAREER HISTORY"),
        ("PROFILE", "WORK HISTORY"),
        ("SUMMARY", "EXPERIENCE"),
    ]
    for index, (left, right) in enumerate(aliases):
        path = output / f"unseen-layout-{index + 1}.pdf"
        _two_column_pdf(
            path,
            name=f"{generator.choice(first_names)} {generator.choice(last_names)}",
            left_heading=left,
            right_heading=right,
            split=260 + index * 18,
        )
        generated.append(path)
    return generated


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("runtime/generalization-validation"),
    )
    parser.add_argument("--seed", type=int, default=90210)
    args = parser.parse_args()
    for path in generate(args.output.resolve(), args.seed):
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
