"""Prepare the supplied resume DOCX files as flow-safe docxtpl templates.

The script always reads retained source backups and writes new files. It never
edits a source template in place.
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

NAVY = RGBColor(0x1A, 0x23, 0x7E)
INK = RGBColor(0x0D, 0x0D, 0x0D)
MUTED = RGBColor(0x59, 0x59, 0x59)


def _set_run_font(
    run: Run,
    family: str,
    size: float,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor = INK,
) -> None:
    run.font.name = family
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), family)


def _clear_body(document: Document) -> None:
    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def _set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: tuple[int, int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    layout = table_properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def _style(document: Document, name: str, family: str, size: float, **kwargs):
    styles = document.styles
    style = styles[name] if name in styles else styles.add_style(name, 1)
    style.font.name = family
    style.font.size = Pt(size)
    style.font.bold = kwargs.get("bold", False)
    style.font.italic = kwargs.get("italic", False)
    style.font.color.rgb = kwargs.get("color", INK)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), family)
    style.paragraph_format.space_before = Pt(kwargs.get("before", 0))
    style.paragraph_format.space_after = Pt(kwargs.get("after", 3))
    style.paragraph_format.line_spacing = kwargs.get("line_spacing", 1.05)
    style.paragraph_format.keep_with_next = kwargs.get("keep_with_next", False)
    return style


def _configure_template_one_styles(document: Document) -> None:
    _style(document, "T1 Name", "Tahoma", 24, bold=True, after=0)
    _style(document, "T1 Title", "Calibri", 11, color=MUTED, after=9)
    _style(
        document,
        "T1 Heading",
        "Tahoma",
        12,
        bold=True,
        color=NAVY,
        before=6,
        after=4,
        keep_with_next=True,
    )
    _style(document, "T1 Entry", "Calibri", 10.5, bold=True, after=1, keep_with_next=True)
    _style(document, "T1 Meta", "Calibri", 9.5, color=MUTED, after=3)
    _style(document, "T1 Body", "Calibri", 10.5, after=3, line_spacing=1.08)
    _style(document, "T1 Sidebar", "Calibri", 10, after=2, line_spacing=1.05)
    bullet = _style(document, "T1 Bullet", "Calibri", 10, after=2, line_spacing=1.05)
    bullet.base_style = document.styles["List Paragraph"]


def _apply_numbering(paragraph: Paragraph, num_id: int = 1) -> None:
    properties = paragraph._p.get_or_add_pPr()
    numbering = properties.find(qn("w:numPr"))
    if numbering is None:
        numbering = OxmlElement("w:numPr")
        properties.append(numbering)
    level = numbering.find(qn("w:ilvl"))
    if level is None:
        level = OxmlElement("w:ilvl")
        numbering.append(level)
    level.set(qn("w:val"), "0")
    identifier = numbering.find(qn("w:numId"))
    if identifier is None:
        identifier = OxmlElement("w:numId")
        numbering.append(identifier)
    identifier.set(qn("w:val"), str(num_id))
    indentation = properties.find(qn("w:ind"))
    if indentation is None:
        indentation = OxmlElement("w:ind")
        properties.append(indentation)
    indentation.set(qn("w:left"), "360")
    indentation.set(qn("w:hanging"), "180")


def _paragraph(container, text: str, style: str, *, reuse_first: bool = False) -> Paragraph:
    if reuse_first and len(container.paragraphs) == 1 and not container.paragraphs[0].text:
        paragraph = container.paragraphs[0]
        paragraph.style = style
    else:
        paragraph = container.add_paragraph(style=style)
    paragraph.add_run(text)
    if style == "T1 Bullet":
        _apply_numbering(paragraph)
    return paragraph


def _tag(container, expression: str, *, reuse_first: bool = False) -> Paragraph:
    return _paragraph(container, expression, "T1 Body", reuse_first=reuse_first)


def _template_one(source: Path, destination: Path) -> None:
    document = Document(source)
    _clear_body(document)
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    _configure_template_one_styles(document)

    table = document.add_table(rows=0, cols=2)
    widths = (6900, 3600)

    def row(*, sidebar_top: int = 0, cant_split: bool = False):
        main_cell, sidebar_cell = table.add_row().cells
        if cant_split:
            row_properties = table.rows[-1]._tr.get_or_add_trPr()
            row_properties.append(OxmlElement("w:cantSplit"))
        main_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        sidebar_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_margins(main_cell, top=0, start=0, bottom=0, end=180)
        _set_cell_margins(sidebar_cell, top=sidebar_top, start=180, bottom=0, end=0)
        return main_cell, sidebar_cell

    def row_tag(expression: str) -> None:
        main_cell, _sidebar_cell = row()
        _paragraph(main_cell, expression, "T1 Body", reuse_first=True)

    def certification_content(container, *, reuse_first: bool = False) -> None:
        _paragraph(container, "Certifications", "T1 Heading", reuse_first=reuse_first)
        _tag(container, "{%p for item in certifications %}")
        _paragraph(container, "{{ item.name }}", "T1 Bullet")
        _tag(container, "{%p if item.issuer or item.date or item.credential_id %}")
        _paragraph(
            container,
            "{% if item.issuer %}{{ item.issuer }}{% endif %}{% if item.date %} | {{ item.date }}{% endif %}{% if item.credential_id %} | {{ item.credential_id }}{% endif %}",
            "T1 Sidebar",
        )
        _tag(container, "{%p endif %}")
        _tag(container, "{%p if item.url %}")
        _paragraph(container, "{{ item.url }}", "T1 Sidebar")
        _tag(container, "{%p endif %}")
        _tag(container, "{%p endfor %}")

    main, sidebar = row(sidebar_top=360)

    name = main.paragraphs[0]
    name.style = "T1 Name"
    first = name.add_run("{{ name_first }}")
    rest = name.add_run("{% if name_rest %} {{ name_rest }}{% endif %}")
    _set_run_font(first, "Tahoma", 24, bold=True, color=NAVY)
    _set_run_font(rest, "Tahoma", 24, bold=True, color=INK)
    _tag(main, "{%p if job_title %}")
    _paragraph(main, "{{ job_title }}", "T1 Title")
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if summary %}")
    _paragraph(main, "Professional Summary", "T1 Heading")
    _paragraph(main, "{{ summary }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(sidebar, "{%p if contact_items %}", reuse_first=True)
    _paragraph(sidebar, "Contact Information", "T1 Heading")
    _tag(sidebar, "{%p for item in contact_items %}")
    _paragraph(sidebar, "{{ item.label }}", "T1 Entry")
    _paragraph(sidebar, "{{ item.value }}", "T1 Sidebar")
    _tag(sidebar, "{%p endfor %}")
    _tag(sidebar, "{%p endif %}")
    _tag(sidebar, "{%p if skills %}")
    _paragraph(sidebar, "Skills", "T1 Heading")
    _tag(sidebar, "{%p for group in skills %}")
    _paragraph(sidebar, "{{ group.group }}", "T1 Entry")
    _paragraph(sidebar, "{{ group['items'] | join(', ') }}", "T1 Sidebar")
    _tag(sidebar, "{%p endfor %}")
    _tag(sidebar, "{%p endif %}")
    _paragraph(sidebar, "", "T1 Sidebar")

    row_tag("{%tr if experience or languages %}")
    main, sidebar = row()
    _tag(main, "{%p if experience %}", reuse_first=True)
    _paragraph(main, "Professional Experience", "T1 Heading")
    _tag(main, "{%p endif %}")
    _tag(sidebar, "{%p if languages %}", reuse_first=True)
    _paragraph(sidebar, "Languages", "T1 Heading")
    _tag(sidebar, "{%p for item in languages %}")
    _paragraph(
        sidebar,
        "{{ item.language }}{% if item.proficiency %}: {{ item.proficiency }}{% endif %}{% if item.cefr %} ({{ item.cefr }}){% endif %}",
        "T1 Bullet",
    )
    _tag(sidebar, "{%p endfor %}")
    _tag(sidebar, "{%p endif %}")
    row_tag("{%tr endif %}")

    row_tag("{%tr for exp in experience %}")
    main, _sidebar = row()
    _paragraph(main, "{{ exp.job_title or '' }}", "T1 Entry", reuse_first=True)
    _paragraph(
        main,
        "{{ exp.organization_line }}{% if exp.date_range %} | {{ exp.date_range }}{% endif %}",
        "T1 Meta",
    )
    _tag(main, "{%p for bullet in exp.bullets %}")
    _paragraph(main, "{{ bullet }}", "T1 Bullet")
    _tag(main, "{%p endfor %}")
    _tag(main, "{%p if exp.technologies %}")
    _paragraph(main, "Technologies: {{ exp.technologies | join(', ') }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if exp.metrics %}")
    _paragraph(main, "Metrics: {{ exp.metrics | join(', ') }}", "T1 Body")
    _tag(main, "{%p endif %}")
    row_tag("{%tr endfor %}")

    row_tag("{%tr for project in projects %}")
    main, sidebar = row(cant_split=True)
    _tag(main, "{%p if loop.first %}", reuse_first=True)
    _paragraph(main, "Projects", "T1 Heading")
    _tag(main, "{%p endif %}")
    _paragraph(main, "{{ project.name or 'Project' }}", "T1 Entry")
    _tag(main, "{%p if project.role or project.date_range %}")
    _paragraph(
        main,
        "{% if project.role %}{{ project.role }}{% endif %}{% if project.date_range %} | {{ project.date_range }}{% endif %}",
        "T1 Meta",
    )
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if project.description %}")
    _paragraph(main, "{{ project.description }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if project.technologies %}")
    _paragraph(main, "Technologies: {{ project.technologies | join(', ') }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if project.url %}")
    _paragraph(main, "{{ project.url }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(sidebar, "{%p if loop.first and certifications %}", reuse_first=True)
    certification_content(sidebar)
    _tag(sidebar, "{%p endif %}")
    _paragraph(sidebar, "", "T1 Sidebar")
    row_tag("{%tr endfor %}")

    row_tag("{%tr if not projects and certifications %}")
    main, sidebar = row(cant_split=True)
    _paragraph(main, "", "T1 Body", reuse_first=True)
    certification_content(sidebar, reuse_first=True)
    row_tag("{%tr endif %}")

    row_tag("{%tr if education %}")
    main, _sidebar = row()
    _paragraph(main, "Education", "T1 Heading", reuse_first=True)
    row_tag("{%tr endif %}")

    row_tag("{%tr for item in education %}")
    main, _sidebar = row()
    _paragraph(main, "{{ item.degree_line or 'Education' }}", "T1 Entry", reuse_first=True)
    _paragraph(
        main,
        "{{ item.institution_line }}{% if item.date_range %} | {{ item.date_range }}{% endif %}",
        "T1 Meta",
    )
    _tag(main, "{%p if item.description %}")
    _paragraph(main, "{{ item.description }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if item.gpa %}")
    _paragraph(main, "GPA: {{ item.gpa }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if item.honors %}")
    _paragraph(main, "Honors: {{ item.honors | join(', ') }}", "T1 Body")
    _tag(main, "{%p endif %}")
    _tag(main, "{%p if item.coursework %}")
    _paragraph(main, "Coursework: {{ item.coursework | join(', ') }}", "T1 Body")
    _tag(main, "{%p endif %}")
    row_tag("{%tr endfor %}")

    _set_table_geometry(table, widths)
    document.save(destination)


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    run_properties = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            run_properties = deepcopy(run._r.rPr)
            break
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run_element = OxmlElement("w:r")
    if run_properties is not None:
        run_element.append(run_properties)
    paragraph._p.append(run_element)
    Run(run_element, paragraph).text = text


def _clone_paragraph(source: Paragraph, text: str) -> Paragraph:
    element = deepcopy(source._p)
    paragraph = Paragraph(element, source._parent)
    _replace_paragraph_text(paragraph, text)
    return paragraph


def _append(document: Document, source: Paragraph, text: str) -> Paragraph:
    paragraph = _clone_paragraph(source, text)
    document._element.body.sectPr.addprevious(paragraph._p)
    return paragraph


def _template_two(source: Path, destination: Path) -> None:
    document = Document(source)
    paragraphs = list(document.paragraphs)
    samples = {
        "name": paragraphs[0],
        "title": paragraphs[1],
        "contact": paragraphs[2],
        "blank": paragraphs[3],
        "heading": paragraphs[4],
        "body": paragraphs[5],
        "entry": paragraphs[8],
        "meta": paragraphs[9],
        "bullet": paragraphs[10],
    }
    _clear_body(document)

    def add(kind: str, text: str) -> Paragraph:
        paragraph = _append(document, samples[kind], text)
        if kind in {"heading", "entry"}:
            paragraph.paragraph_format.keep_with_next = True
        if kind == "meta":
            for run in paragraph.runs:
                run.font.italic = False
                fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                    fonts.set(qn(f"w:{attribute}"), "Arial")
        return paragraph

    def tag(text: str) -> None:
        add("blank", text)

    add("name", "{{ name }}")
    tag("{%p if job_title %}")
    add("title", "{{ job_title }}")
    tag("{%p endif %}")
    tag("{%p if contact_line %}")
    add("contact", "{{ contact_line }}")
    tag("{%p endif %}")
    add("blank", "")

    tag("{%p if summary %}")
    add("heading", "PROFESSIONAL SUMMARY")
    add("body", "{{ summary }}")
    add("blank", "")
    tag("{%p endif %}")

    tag("{%p if experience %}")
    add("heading", "WORK EXPERIENCE")
    tag("{%p for exp in experience %}")
    add("entry", "{{ exp.job_title or '' }}")
    add("meta", "{{ exp.organization_line }}\t{{ exp.date_range }}")
    tag("{%p for bullet in exp.bullets %}")
    add("bullet", "{{ bullet }}")
    tag("{%p endfor %}")
    tag("{%p if exp.technologies %}")
    add("body", "Technologies: {{ exp.technologies | join(', ') }}")
    tag("{%p endif %}")
    tag("{%p if exp.metrics %}")
    add("body", "Metrics: {{ exp.metrics | join(', ') }}")
    tag("{%p endif %}")
    tag("{%p endfor %}")
    add("blank", "")
    tag("{%p endif %}")

    tag("{%p if projects %}")
    add("heading", "PROJECTS")
    tag("{%p for project in projects %}")
    add("entry", "{{ project.name or 'Project' }}")
    tag("{%p if project.role or project.date_range %}")
    add("meta", "{{ project.role or '' }}\t{{ project.date_range }}")
    tag("{%p endif %}")
    tag("{%p if project.description %}")
    add("body", "{{ project.description }}")
    tag("{%p endif %}")
    tag("{%p if project.technologies %}")
    add("body", "Technologies: {{ project.technologies | join(', ') }}")
    tag("{%p endif %}")
    tag("{%p if project.url %}")
    add("body", "{{ project.url }}")
    tag("{%p endif %}")
    tag("{%p endfor %}")
    add("blank", "")
    tag("{%p endif %}")

    tag("{%p if education %}")
    add("heading", "EDUCATION")
    tag("{%p for item in education %}")
    add("entry", "{{ item.degree_line or 'Education' }}")
    add("meta", "{{ item.institution_line }}\t{{ item.date_range }}")
    tag("{%p if item.description %}")
    add("body", "{{ item.description }}")
    tag("{%p endif %}")
    tag("{%p if item.gpa %}")
    add("body", "GPA: {{ item.gpa }}")
    tag("{%p endif %}")
    tag("{%p if item.honors %}")
    add("body", "Honors: {{ item.honors | join(', ') }}")
    tag("{%p endif %}")
    tag("{%p if item.coursework %}")
    add("body", "Coursework: {{ item.coursework | join(', ') }}")
    tag("{%p endif %}")
    tag("{%p endfor %}")
    add("blank", "")
    tag("{%p endif %}")

    tag("{%p if skills %}")
    add("heading", "SKILLS")
    tag("{%p for group in skills %}")
    add("entry", "{{ group.group }}")
    tag("{%p for skill in group['items'] %}")
    add("bullet", "{{ skill }}")
    tag("{%p endfor %}")
    tag("{%p endfor %}")
    add("blank", "")
    tag("{%p endif %}")

    tag("{%p if languages %}")
    add("heading", "LANGUAGES")
    tag("{%p for item in languages %}")
    add(
        "bullet",
        "{{ item.language }}{% if item.proficiency %}: {{ item.proficiency }}{% endif %}{% if item.cefr %} ({{ item.cefr }}){% endif %}",
    )
    tag("{%p endfor %}")
    add("blank", "")
    tag("{%p endif %}")

    tag("{%p if certifications %}")
    add("heading", "CERTIFICATIONS")
    tag("{%p for item in certifications %}")
    add("bullet", "{{ item.name }}")
    tag("{%p if item.issuer or item.date or item.credential_id %}")
    add(
        "body",
        "{% if item.issuer %}{{ item.issuer }}{% endif %}{% if item.date %} | {{ item.date }}{% endif %}{% if item.credential_id %} | {{ item.credential_id }}{% endif %}",
    )
    tag("{%p endif %}")
    tag("{%p if item.url %}")
    add("body", "{{ item.url }}")
    tag("{%p endif %}")
    tag("{%p endfor %}")
    tag("{%p endif %}")
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    _template_one(args.source_dir / "Template-1.docx", args.output_dir / "Template-1.docx")
    _template_two(args.source_dir / "Template-2.docx", args.output_dir / "Template-2.docx")


if __name__ == "__main__":
    main()
