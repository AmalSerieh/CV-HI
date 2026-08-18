"""Privacy-safe analysis and capability API."""

from __future__ import annotations

import json
from typing import Annotated
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, JSONResponse, Response
from pydantic import ValidationError

from resume_analyzer.diagnostics.health import application_health, system_capabilities
from resume_analyzer.diagnostics.models import model_status
from resume_analyzer.export import (
    DOCX_MEDIA_TYPE,
    DocxRenderError,
    FinalResumeBuilder,
    ReviewStateError,
    ReviewUpdate,
    TemplateNotFound,
    TemplateSelection,
    content_disposition,
)

from ..build_info import build_state
from ..models import AnalysisOptions
from ..services import AnalysisNotFound, TooManyAnalyses, UploadValidationError

router = APIRouter(prefix="/api")


def _boolean(value: str | None, default: bool) -> bool:
    if value is None:
        return default
    normalized = value.strip().casefold()
    if normalized in {"1", "true", "yes", "on"}:
        return True
    if normalized in {"0", "false", "no", "off"}:
        return False
    raise HTTPException(status_code=422, detail="A feature toggle has an invalid value.")


def _identifier(value: str) -> UUID:
    try:
        return UUID(value)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail="Invalid analysis identifier.") from exc


def _record(request: Request, analysis_id: str):
    try:
        return request.app.state.job_store.get(_identifier(analysis_id))
    except AnalysisNotFound as exc:
        raise HTTPException(status_code=404, detail="Analysis not found or expired.") from exc


def _completed_record(request: Request, analysis_id: str):
    record = _record(request, analysis_id)
    if record.status != "completed" or record.result is None:
        raise HTTPException(status_code=409, detail="The result is not available yet.")
    return record


def _builder_and_state(request: Request, analysis_id: str):
    record = _completed_record(request, analysis_id)
    try:
        builder = FinalResumeBuilder(record.result)
        state = request.app.state.job_store.get_or_create_review_state(
            record.id, builder.initial_state
        )
    except (ValidationError, ReviewStateError) as exc:
        raise HTTPException(status_code=409, detail="Resume review state is unavailable.") from exc
    return record, builder, state


@router.post("/analyses", status_code=202)
async def create_analysis(
    request: Request,
    resume: Annotated[UploadFile, File()],
    job_description_text: str | None = Form(default=None),
    job_description_file: Annotated[UploadFile | None, File()] = None,
    enable_target_role: str | None = Form(default=None),
    enable_recommendations: str | None = Form(default=None),
    enable_ats: str | None = Form(default=None),
    enable_job_match: str | None = Form(default=None),
    enable_rewrites: str | None = Form(default=None),
    enable_ocr: str | None = Form(default=None),
    ai_provider: str = Form(default="none"),
    ai_model: str | None = Form(default=None),
    output_language: str = Form(default="auto"),
    rewrite_summary: str | None = Form(default=None),
    rewrite_experience: str | None = Form(default=None),
    rewrite_skills: str | None = Form(default=None),
    bullet_rewrite_mode: str = Form(default="first"),
    bullet_rewrite_count: int = Form(default=20),
    bullet_rewrite_selection: str = Form(default=""),
):
    source = build_state(
        request.app.state.startup_source_fingerprint,
        request.app.state.source_fingerprint_provider,
    )
    if source["restart_required"]:
        return JSONResponse(
            status_code=503,
            content={
                "error": {
                    "code": "backend_restart_required",
                    "message": (
                        "The backend source changed after this server started. "
                        "Restart the local web application before analyzing another resume."
                    ),
                },
                "build": source,
            },
            headers={"Retry-After": "1"},
        )
    uploads = request.app.state.upload_service
    try:
        prepared = await uploads.prepare(
            resume,
            job_description_text=job_description_text,
            job_description_file=job_description_file,
        )
    except UploadValidationError as exc:
        return JSONResponse(
            status_code=exc.status_code,
            content={"error": {"code": exc.code, "message": exc.message}},
        )
    language = None if output_language == "auto" else output_language
    if language not in {None, "en", "ar"}:
        uploads.cleanup(prepared.directory)
        raise HTTPException(status_code=422, detail="Invalid output language.")
    sections = tuple(
        name
        for name, enabled in (
            ("summary", _boolean(rewrite_summary, True)),
            ("experience", _boolean(rewrite_experience, True)),
            ("skills", _boolean(rewrite_skills, True)),
        )
        if enabled
    )
    if not sections:
        uploads.cleanup(prepared.directory)
        raise HTTPException(status_code=422, detail="Select at least one rewrite section.")
    bullet_rewrite_mode = bullet_rewrite_mode.strip().casefold()
    if bullet_rewrite_mode not in {"first", "specific", "all"}:
        uploads.cleanup(prepared.directory)
        raise HTTPException(status_code=422, detail="Invalid bullet rewrite mode.")
    if not 1 <= bullet_rewrite_count <= 20:
        uploads.cleanup(prepared.directory)
        raise HTTPException(
            status_code=422, detail="Bullet rewrite count must be between 1 and 20."
        )
    try:
        selected_bullets = tuple(
            dict.fromkeys(
                int(value.strip()) - 1
                for value in bullet_rewrite_selection.split(",")
                if value.strip()
            )
        )
    except ValueError as exc:
        uploads.cleanup(prepared.directory)
        raise HTTPException(
            status_code=422, detail="Specific bullet numbers must be comma-separated integers."
        ) from exc
    if bullet_rewrite_mode == "specific" and (
        not selected_bullets
        or any(value < 0 for value in selected_bullets)
        or len(selected_bullets) > 20
    ):
        uploads.cleanup(prepared.directory)
        raise HTTPException(
            status_code=422, detail="Select between 1 and 20 positive bullet numbers."
        )
    options = AnalysisOptions(
        enable_target_role=_boolean(enable_target_role, True),
        enable_recommendations=_boolean(enable_recommendations, True),
        enable_ats=_boolean(enable_ats, True),
        enable_job_match=_boolean(enable_job_match, True),
        enable_rewrites=_boolean(enable_rewrites, False),
        enable_ocr=_boolean(enable_ocr, True),
        ai_provider=ai_provider,
        ai_model=ai_model,
        output_language=language,
        rewrite_sections=sections,
        bullet_rewrite_mode=bullet_rewrite_mode,
        bullet_rewrite_count=bullet_rewrite_count,
        bullet_rewrite_selection=(selected_bullets if bullet_rewrite_mode == "specific" else None),
    )
    try:
        record = request.app.state.analysis_service.submit(prepared, options)
    except TooManyAnalyses:
        uploads.cleanup(prepared.directory)
        return JSONResponse(
            status_code=429,
            content={
                "error": {
                    "code": "analysis_capacity_reached",
                    "message": "The local analysis capacity is full. Try again shortly.",
                }
            },
            headers={"Retry-After": "5"},
        )
    return {
        "analysis_id": str(record.id),
        "status": record.status,
        "status_url": f"/api/analyses/{record.id}",
        "result_url": f"/api/analyses/{record.id}/result",
        "page_url": f"/results/{record.id}",
    }


@router.get("/analyses/{analysis_id}")
def analysis_status(request: Request, analysis_id: str):
    return _record(request, analysis_id).public_status()


@router.get("/analyses/{analysis_id}/result")
def analysis_result(request: Request, analysis_id: str):
    record = _completed_record(request, analysis_id)
    return record.result


@router.get("/analyses/{analysis_id}/download")
def analysis_download(request: Request, analysis_id: str):
    record = _completed_record(request, analysis_id)
    payload = json.dumps(record.result, ensure_ascii=False, indent=2, allow_nan=False) + "\n"
    return Response(
        content=payload.encode("utf-8"),
        media_type="application/json; charset=utf-8",
        headers={
            "Content-Disposition": f'attachment; filename="analysis-{record.id}.json"',
            "X-Content-Type-Options": "nosniff",
        },
    )


@router.get("/analyses/{analysis_id}/resume-review")
def resume_review(request: Request, analysis_id: str):
    _record_value, builder, state = _builder_and_state(request, analysis_id)
    return builder.review_payload(state).model_dump(mode="json")


@router.patch("/analyses/{analysis_id}/resume-review")
def update_resume_review(request: Request, analysis_id: str, update: ReviewUpdate):
    record = _completed_record(request, analysis_id)
    try:
        builder = FinalResumeBuilder(record.result)
        state = request.app.state.job_store.update_review_state(
            record.id,
            builder.initial_state,
            lambda current: builder.apply_update(current, update),
        )
    except ReviewStateError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except ValidationError as exc:
        raise HTTPException(status_code=409, detail="Resume review state is unavailable.") from exc
    return builder.review_payload(state).model_dump(mode="json")


@router.get("/analyses/{analysis_id}/final-resume")
def final_resume(request: Request, analysis_id: str):
    _record_value, builder, state = _builder_and_state(request, analysis_id)
    return builder.build(state).model_dump(mode="json")


@router.get("/resume-templates")
def resume_templates(request: Request):
    return request.app.state.template_registry.public_metadata()


@router.get("/resume-templates/{template_id}/preview", response_class=FileResponse)
def resume_template_preview(request: Request, template_id: str):
    try:
        definition = request.app.state.template_registry.get(template_id)
    except TemplateNotFound as exc:
        raise HTTPException(status_code=404, detail="Resume template not found.") from exc
    if not definition.preview_path.is_file():
        raise HTTPException(status_code=404, detail="Resume template preview is unavailable.")
    media_type = (
        "image/jpeg" if definition.preview_path.suffix.casefold() in {".jpg", ".jpeg"} else None
    )
    return FileResponse(definition.preview_path, media_type=media_type)


@router.post("/analyses/{analysis_id}/download-docx")
def download_resume_docx(request: Request, analysis_id: str, selection: TemplateSelection):
    _record_value, builder, state = _builder_and_state(request, analysis_id)
    final = builder.build(state)
    try:
        content = request.app.state.docx_renderer.render(final, selection.template_id)
    except TemplateNotFound as exc:
        raise HTTPException(status_code=422, detail="Unsupported resume template.") from exc
    except DocxRenderError as exc:
        raise HTTPException(
            status_code=503, detail="The Word resume could not be generated."
        ) from exc
    return Response(
        content=content,
        media_type=DOCX_MEDIA_TYPE,
        headers={
            "Content-Disposition": content_disposition(final.contact.name),
            "X-Content-Type-Options": "nosniff",
        },
    )


@router.delete("/analyses/{analysis_id}", status_code=204)
def delete_analysis(request: Request, analysis_id: str):
    try:
        request.app.state.job_store.delete(_identifier(analysis_id))
    except AnalysisNotFound as exc:
        raise HTTPException(status_code=404, detail="Analysis not found or expired.") from exc
    return Response(status_code=204)


@router.get("/health")
def health(request: Request):
    payload = application_health(request.app.state.settings)
    payload["build"] = build_state(
        request.app.state.startup_source_fingerprint,
        request.app.state.source_fingerprint_provider,
    )
    return payload


@router.get("/system")
def system(request: Request):
    return system_capabilities(request.app.state.settings, public=True)


@router.get("/models")
def models():
    return model_status(public=True)

@router.get("/analyses/{analysis_id}/download-docxx")
def download_docx(request: Request, analysis_id: str, template: str = "single_column"):
    record = _record(request, analysis_id)
    if record.status != "completed" or record.result is None:
        raise HTTPException(status_code=409, detail="The result is not available yet.")

    report = record.result

    # 1. استخراج البيانات والترتيب الديناميكي
    entities = report.get("entities") or {}
    contact = entities.get("contact") or {}
    extraction = report.get("extraction") or {}
    raw_sections = extraction.get("sections") or {}
    section_order = extraction.get("section_order") or list(raw_sections.keys())
    rewrites = report.get("rewrites") or {}

    # تجميع كافة مفاتيح الأقسام المتاحة
    all_section_keys = []
    for k in section_order + list(raw_sections.keys()) + list(entities.keys()):
        if k not in all_section_keys and k not in ("contact", "contact_header"):
            all_section_keys.append(k)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    def get_heading(sec_key):
        sec_data = raw_sections.get(sec_key) or {}
        return sec_data.get("heading") or sec_key.replace("_", " ").title()

    # معلومات التواصل والاسم
    name = contact.get("name") or (report.get("document") or {}).get("name", "Candidate Name")
    target_role = report.get("target_role") or {}
    primary_role = target_role.get("primary") or {}
    job_title = contact.get("job_title") or primary_role.get("title_en", "")

    # ================= 1. نموذج العمود الواحد (Classic) =================
    if template == "single_column":
        # الترويسة
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_name = p_name.add_run(str(name).upper())
        r_name.bold = True
        r_name.font.size = Pt(18)

        if job_title:
            p_job = doc.add_paragraph()
            p_job.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_job = p_job.add_run(str(job_title))
            r_job.font.size = Pt(11)
            r_job.font.color.rgb = RGBColor(80, 80, 80)

        details = [v for v in [contact.get("phone"), contact.get("location"), contact.get("email"), contact.get("linkedin")] if v]
        if details:
            p_det = doc.add_paragraph()
            p_det.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_det.add_run("  |  ".join(details)).font.size = Pt(9.5)

        # الأقسام الديناميكية
        for sec_key in all_section_keys:
            heading = get_heading(sec_key)
            raw_content = (raw_sections.get(sec_key) or {}).get("content", "").strip()

            p_h = doc.add_heading(level=2)
            p_h.add_run(heading.upper()).bold = True
            p_h.paragraph_format.space_before = Pt(14)
            p_h.paragraph_format.space_after = Pt(4)

            if sec_key == "summary":
                text = (rewrites.get("summary") or {}).get("improved") or raw_content or entities.get("summary", "")
                if text: doc.add_paragraph(text)

            elif sec_key == "experience":
                exps = entities.get("experience") or []
                exp_rewrites = rewrites.get("experience_bullets") or []

                if exps:
                    for exp in exps:
                        t = exp.get("job_title") or "Position"
                        c = exp.get("company") or ""
                        dates = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}" if exp.get("start_date") else ""

                        table = doc.add_table(rows=1, cols=2)
                        table.autofit = False
                        table.columns[0].width = Inches(5.2)
                        table.columns[1].width = Inches(1.3)
                        row = table.rows[0]
                        row.cells[0].paragraphs[0].add_run(f"{t} | {c}" if c else t).bold = True
                        p_r = row.cells[1].paragraphs[0]
                        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p_r.add_run(dates)

                        for b in exp.get("responsibilities") or exp.get("bullets") or []:
                            b_str = b.get("text", "") if isinstance(b, dict) else str(b)
                            final_text = b_str
                            for r in exp_rewrites:
                                if r.get("original", "").strip() == b_str.strip() and r.get("improved"):
                                    final_text = r.get("improved")
                                    break
                            doc.add_paragraph(final_text.lstrip("•- "), style='List Bullet')
                elif raw_content:
                    for line in raw_content.split("\n"):
                        if line.strip(): doc.add_paragraph(line.lstrip("•- ").strip())

            else:
                if raw_content:
                    for line in raw_content.split("\n"):
                        if line.strip():
                            is_bullet = line.strip().startswith("•") or line.strip().startswith("-")
                            doc.add_paragraph(line.lstrip("•- ").strip(), style='List Bullet' if is_bullet else 'Normal')

    # ================= 2. نموذج العمودين والشريط الجانبي (Modern Two-Column) =================
    else:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(4.8) # 68% العمود الرئيسي
        table.columns[1].width = Inches(2.2) # 32% الشريط الجانبي
        row = table.rows[0]
        cell_left = row.cells[0]
        cell_right = row.cells[1]

        # اليسار: الاسم والخبرات والملخص
        p_l_name = cell_left.paragraphs[0]
        p_l_name.add_run(str(name)).bold = True
        p_l_name.runs[0].font.size = Pt(16)
        p_l_name.runs[0].font.color.rgb = RGBColor(15, 41, 66)

        if job_title:
            cell_left.add_paragraph(job_title).runs[0].font.color.rgb = RGBColor(100, 100, 100)

        sidebar_keys = {"skills", "languages", "certifications", "contact"}

        for sec_key in all_section_keys:
            if sec_key in sidebar_keys: continue
            heading = get_heading(sec_key)
            p_h = cell_left.add_paragraph()
            r_h = p_h.add_run(heading.upper())
            r_h.bold = True
            r_h.font.color.rgb = RGBColor(15, 41, 66)

            raw_content = (raw_sections.get(sec_key) or {}).get("content", "").strip()
            if raw_content:
                for line in raw_content.split("\n"):
                    if line.strip():
                        cell_left.add_paragraph(line.lstrip("•- ").strip())

        # اليمين: معلومات التواصل والمهارات واللغات
        p_r_head = cell_right.paragraphs[0]
        p_r_head.add_run("CONTACT").bold = True
        p_r_head.runs[0].font.color.rgb = RGBColor(15, 41, 66)

        for detail in [contact.get("phone"), contact.get("email"), contact.get("location"), contact.get("linkedin")]:
            if detail: cell_right.add_paragraph(detail)

        for sec_key in all_section_keys:
            if sec_key not in sidebar_keys: continue
            heading = get_heading(sec_key)
            p_h = cell_right.add_paragraph()
            r_h = p_h.add_run(heading.upper())
            r_h.bold = True
            r_h.font.color.rgb = RGBColor(15, 41, 66)

            raw_content = (raw_sections.get(sec_key) or {}).get("content", "").strip()
            if raw_content:
                for line in raw_content.split("\n"):
                    if line.strip(): cell_right.add_paragraph(line.lstrip("•- ").strip())

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = f"Optimized-Resume-{template}-{record.id}.docx"
    return Response(
        content=stream.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

